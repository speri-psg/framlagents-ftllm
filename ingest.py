# ingest.py — Run once to build the local knowledge base
# Usage: python ingest.py
# Re-run whenever you add/update documents in the docs/ folder

import chromadb
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
import pypdf
import docx
import pandas as pd
import os
import glob

DOCS_FOLDER = "docs"
CHROMA_PATH = "chroma_db"
CHUNK_SIZE = 150    # words per chunk — all-MiniLM-L6-v2 max is 256 tokens (~180 words)
CHUNK_OVERLAP = 20  # words overlap between chunks


_FRENCH_MARKERS = ("l'Agence", "la criminalité", "du financement", "les activités terroristes",
                   "Recyclage des produits", "le blanchiment", "conformément", "paragraphe")


def _is_mostly_french(text: str) -> bool:
    """Return True if chunk appears to be French (bilingual PDF noise)."""
    hits = sum(1 for m in _FRENCH_MARKERS if m in text)
    return hits >= 2


def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    words = text.split()
    chunks = []
    for i in range(0, len(words), size - overlap):
        chunk = " ".join(words[i:i + size])
        if chunk.strip() and not _is_mostly_french(chunk):
            chunks.append(chunk)
    return chunks


def ingest_pdf(path):
    reader = pypdf.PdfReader(path)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text)
    full_text = " ".join(pages_text)
    return chunk_text(full_text)


def ingest_docx(path):
    try:
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        full_text = " ".join(parts)
    except Exception:
        # Fallback: extract raw text from word/document.xml directly
        import zipfile, re
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        full_text = re.sub(r"<[^>]+>", " ", xml)
        full_text = re.sub(r"\s+", " ", full_text).strip()
    return chunk_text(full_text)


def ingest_csv(path):
    df = pd.read_csv(path)
    chunks = [df[i:i + 30].to_string() for i in range(0, len(df), 30)]
    return [c for c in chunks if c.strip()]


def main():
    ef = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
    client = chromadb.PersistentClient(path=CHROMA_PATH)

    # Fresh ingest — drop existing collection
    try:
        client.delete_collection("framl_kb")
        print("Dropped existing collection for fresh ingest.")
    except Exception:
        pass

    collection = client.create_collection("framl_kb", embedding_function=ef)

    all_chunks, all_ids, all_metadata = [], [], []
    chunk_id = 0

    # Ingest PDFs from docs/
    pdf_files = glob.glob(os.path.join(DOCS_FOLDER, "*.pdf"))
    if not pdf_files:
        print(f"No PDFs found in {DOCS_FOLDER}/  — add your documents there and re-run.")
    for pdf_path in pdf_files:
        fname = os.path.basename(pdf_path)
        print(f"Ingesting PDF: {fname} ...")
        try:
            chunks = ingest_pdf(pdf_path)
            for chunk in chunks:
                all_chunks.append(chunk)
                all_ids.append(f"chunk_{chunk_id}")
                all_metadata.append({"source": fname, "type": "pdf"})
                chunk_id += 1
            print(f"  -> {len(chunks)} chunks")
        except Exception as e:
            print(f"  ERROR reading {fname}: {e}")

    # Ingest plain text files from docs/
    txt_files = glob.glob(os.path.join(DOCS_FOLDER, "*.txt"))
    for txt_path in txt_files:
        fname = os.path.basename(txt_path)
        print(f"Ingesting text file: {fname} ...")
        try:
            with open(txt_path, encoding="utf-8") as f:
                full_text = f.read()
            chunks = chunk_text(full_text)
            for chunk in chunks:
                all_chunks.append(chunk)
                all_ids.append(f"chunk_{chunk_id}")
                all_metadata.append({"source": fname, "type": "txt"})
                chunk_id += 1
            print(f"  -> {len(chunks)} chunks")
        except Exception as e:
            print(f"  ERROR reading {fname}: {e}")

    # Ingest Word docs from docs/
    docx_files = glob.glob(os.path.join(DOCS_FOLDER, "*.docx"))
    for docx_path in docx_files:
        fname = os.path.basename(docx_path)
        print(f"Ingesting Word doc: {fname} ...")
        try:
            chunks = ingest_docx(docx_path)
            for chunk in chunks:
                all_chunks.append(chunk)
                all_ids.append(f"chunk_{chunk_id}")
                all_metadata.append({"source": fname, "type": "docx"})
                chunk_id += 1
            print(f"  -> {len(chunks)} chunks")
        except Exception as e:
            print(f"  ERROR reading {fname}: {e}")

    # CSVs are NOT ingested into ChromaDB — data analysis is done directly
    # via pandas in the app. ChromaDB is reserved for policy/guideline PDFs.

    if not all_chunks:
        print("\nNo documents ingested. Add PDFs/CSVs to the docs/ folder and re-run.")
        return

    # Batch insert into ChromaDB
    BATCH = 100
    print(f"\nEmbedding and storing {chunk_id} chunks ...")
    for i in range(0, len(all_chunks), BATCH):
        collection.add(
            documents=all_chunks[i:i + BATCH],
            ids=all_ids[i:i + BATCH],
            metadatas=all_metadata[i:i + BATCH]
        )
        print(f"  Stored {min(i + BATCH, chunk_id)}/{chunk_id} chunks", end="\r")

    print(f"\nDone. {chunk_id} chunks stored in {CHROMA_PATH}/")
    print("You can now start the app: python application.py")


if __name__ == "__main__":
    main()