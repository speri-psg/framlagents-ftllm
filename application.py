"""
application.py — ARIA (Qwen2.5 via Ollama)

Run:
    python application.py        # http://127.0.0.1:5000

Override the LLM endpoint if needed:
    set OLLAMA_BASE_URL=http://localhost:11434/v1
    set OLLAMA_MODEL=qwen2.5:7b
"""

import io
import base64
import datetime
import json
import re
import time
import sys
import os
import threading

# Force line-buffered stdout so print() appears immediately in any terminal/shell
sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

# ── Ensure project root is on path so config + analytics imports work ─────────
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

import pandas as pd
import flask
import dash_bootstrap_components as dbc
from dash import Dash, callback, html, dcc, Input, Output, State, callback_context, ALL, no_update
from dash import dash_table
from dash_chat import ChatComponent

from config import ALERTS_CSV, DS_CSV, SAR_CSV, CLUSTER_LABELS_CSV, OLLAMA_MODEL, OLLAMA_BASE_URL
from agents import OrchestratorAgent
from agents.base_agent import stop_event as _agent_stop_event
import upload_kb
import lambda_ds_performance
import lambda_rule_analysis
import lambda_ofac
import make_figures

MAX_SWEEP_ROWS = 10  # max sweep points passed to model (prevents token limit cutoff)

# ── Data loading ──────────────────────────────────────────────────────────────
# Auto-detect separator (production file is tab-separated; synth is comma-separated)
with open(ALERTS_CSV, "r", encoding="utf-8") as _f:
    _first_line = _f.readline()
_sep = "\t" if "\t" in _first_line else ","
_df_raw = pd.read_csv(ALERTS_CSV, sep=_sep)

# Rename production column names to canonical names (no-op if already canonical)
_df_raw = _df_raw.rename(columns={
    "AVG_TRXNS_WEEK":  "avg_num_trxns",
    "AVG_TRXN_AMT":    "avg_trxn_amt",
    "TRXN_AMT_MONTHLY":"trxn_amt_monthly",
    "FP":              "false_positives",
    "FN":              "false_negatives",
    "ALERT":           "alerts",
    "CUSTOMER_TYPE":   "customer_type",
})
# Production uses Yes/No strings; synth uses 0/1 integers — normalise to int
for _col in ("alerts", "false_positives", "false_negatives"):
    if _col in _df_raw.columns and _df_raw[_col].dtype == object:
        _df_raw[_col] = _df_raw[_col].map({"Yes": 1, "No": 0})
    elif _col not in _df_raw.columns:
        _df_raw[_col] = 0
if "dynamic_segment" not in _df_raw.columns:
    _df_raw["dynamic_segment"] = _df_raw["customer_type"].map({"BUSINESS": 0, "INDIVIDUAL": 1})
DF           = _df_raw
DF_BUSINESS  = DF[DF["dynamic_segment"] == 0]
DF_INDIVIDUAL= DF[DF["dynamic_segment"] == 1]

DF_SS  = pd.read_csv(DS_CSV) if os.path.exists(DS_CSV) else None
if DF_SS is not None and "dynamic_segment" not in DF_SS.columns:
    DF_SS["dynamic_segment"] = DF_SS["customer_type"].str.upper().map({"BUSINESS": 0, "INDIVIDUAL": 1})
DF_SAR = pd.read_csv(SAR_CSV) if os.path.exists(SAR_CSV) else None
if DF_SAR is not None and "dynamic_segment" not in DF_SAR.columns:
    DF_SAR["dynamic_segment"] = DF_SAR["customer_type"].map({"BUSINESS": 0, "INDIVIDUAL": 1})

# ── Discover segmentation dimensions from data (≥70% availability, ≤20 unique values) ──
import lambda_ds_performance as _ldp_cfg
_DS_DIMS: dict = {}
if DF_SS is not None:
    try:
        from column_mapper import normalize_columns as _norm
        _df_norm = _norm(DF_SS, verbose=False)
        _DS_DIMS = {
            "BUSINESS":   _ldp_cfg.discover_dims(_df_norm, segment="BUSINESS"),
            "INDIVIDUAL": _ldp_cfg.discover_dims(_df_norm, segment="INDIVIDUAL"),
        }
        print(f"Segmentation dims — Business:   {_DS_DIMS['BUSINESS']}")
        print(f"Segmentation dims — Individual: {_DS_DIMS['INDIVIDUAL']}")
    except Exception as _e:
        print(f"discover_dims failed (non-fatal): {_e}")
        _DS_DIMS = {
            "BUSINESS":   ["ACCOUNT_TYPE", "ACCOUNT_AGE_CATEGORY"],
            "INDIVIDUAL": ["ACCOUNT_TYPE", "GENDER", "AGE_CATEGORY", "INCOME_BAND"],
        }

_TXN_CSV = os.path.join(_HERE, "docs", "aml_transactions.csv")
DF_TXN   = pd.read_csv(_TXN_CSV, parse_dates=["txn_date"]) if os.path.exists(_TXN_CSV) else None
_NET_CSV = os.path.join(_HERE, "docs", "network_features.csv")
DF_NET   = pd.read_csv(_NET_CSV).set_index("customer_id") if os.path.exists(_NET_CSV) else None
print(f"Transaction data: {'loaded (' + str(len(DF_TXN)) + ' rows)' if DF_TXN is not None else 'not found — run generate_aml_transactions.py'}")
DF_SAR_BUSINESS    = DF_SAR[DF_SAR["dynamic_segment"] == 0] if DF_SAR is not None else None
DF_SAR_INDIVIDUAL  = DF_SAR[DF_SAR["dynamic_segment"] == 1] if DF_SAR is not None else None

_total      = len(DF)
_biz_count  = len(DF_BUSINESS)
_ind_count  = len(DF_INDIVIDUAL)
_alert_count= int(pd.to_numeric(DF["alerts"],        errors="coerce").sum())
_fp_count   = int(pd.to_numeric(DF["false_positives"], errors="coerce").sum())
print(f"Alerts data: {_total:,} rows | Business={_biz_count:,} Individual={_ind_count:,}")
print(f"SS data: {'loaded (' + str(len(DF_SS)) + ' rows)' if DF_SS is not None else 'not found — run python ds_data_prep.py'}")
_sar_status = (f"loaded ({len(DF_SAR)} rows, {int(DF_SAR['is_sar'].sum())} SARs)" if DF_SAR is not None else "not found - run python simulate_sars.py")
print(f"SAR simulation: {_sar_status}")

# ── SAR propensity model ──────────────────────────────────────────────────────
import sar_scorer as _sar_scorer
_sar_scores_df = None
_sar_roc_auc   = None
if DF_SAR is not None:
    try:
        _, _sar_roc_auc = _sar_scorer.train(DF_SAR)
        _sar_scores_df  = _sar_scorer.score_alerts(DF_SAR)
    except Exception as _e:
        print(f"[sar_scorer] Failed (non-fatal): {_e}")

DF_RULE_SWEEP = lambda_rule_analysis.load_rule_sweep_data()
_rule_status = (
    f"loaded ({len(DF_RULE_SWEEP)} rows, {DF_RULE_SWEEP['risk_factor'].nunique()} rules)"
    if DF_RULE_SWEEP is not None
    else "not found — run python prepare_rule_sweep_data.py"
)
print(f"Rule sweep data: {_rule_status}")

DF_CLUSTER_LABELS = None
if os.path.exists(CLUSTER_LABELS_CSV):
    DF_CLUSTER_LABELS = pd.read_csv(CLUSTER_LABELS_CSV)
    print(f"Cluster labels: loaded ({len(DF_CLUSTER_LABELS):,} customers)")
else:
    print("Cluster labels: not found — run python prepare_cluster_labels.py")

# ── Cluster enrichment helper ─────────────────────────────────────────────────
def _enrich_cluster_df(df_clustered):
    """Add cluster_label, is_sar, is_fp, is_alerted, is_fn columns to df_clustered."""
    df = df_clustered.copy()
    cluster_vals = sorted(df["cluster"].unique())
    label_map = {v: f"C{i+1}" for i, v in enumerate(cluster_vals)}
    df["cluster_label"] = df["cluster"].map(label_map)
    if DF_RULE_SWEEP is not None and "customer_id" in df.columns:
        sar_map   = DF_RULE_SWEEP.groupby("customer_id")["is_sar"].max()
        alerted   = set(DF_RULE_SWEEP["customer_id"].unique())
        df["is_sar"]     = df["customer_id"].map(sar_map).fillna(0).astype(int)
        df["is_alerted"] = df["customer_id"].isin(alerted).astype(int)
        df["is_fp"]      = ((df["is_alerted"] == 1) & (df["is_sar"] == 0)).astype(int)
        df["is_fn"]      = ((df["is_sar"] == 1)     & (df["is_alerted"] == 0)).astype(int)
    else:
        df["is_sar"] = df["is_alerted"] = df["is_fp"] = df["is_fn"] = 0
    return df


def _filter_treemap_node(node_id, df):
    """Filter enriched df_clustered to the subset matching a treemap node_id."""
    import re as _re
    if not node_id or node_id in ("All", ""):
        return df
    if node_id == "SMALL":
        return df
    if not node_id.startswith("CL__"):
        return df
    inner = node_id[4:]
    m = _re.match(r"^C(\d+):", inner)
    if m:
        c_num = int(m.group(1))
        cluster_vals = sorted(df["cluster"].unique())
        if 1 <= c_num <= len(cluster_vals):
            df = df[df["cluster"] == cluster_vals[c_num - 1]]
    if "__ct_" not in inner:
        return df
    _, ct_rest = inner.split("__ct_", 1)
    ct_parts = ct_rest.split("__", 1)
    ct = ct_parts[0]
    if ct and "customer_type" in df.columns:
        df = df[df["customer_type"] == ct]
    if len(ct_parts) < 2:
        return df
    # Handle multiple dim levels: ACCOUNT_TYPE_Other__GENDER_Male__AGE_CATEGORY_Senior...
    for dim_val in ct_parts[1].split("__"):
        if not dim_val:
            continue
        for col in sorted(df.columns, key=len, reverse=True):
            prefix = col + "_"
            if dim_val.startswith(prefix):
                val = dim_val[len(prefix):]
                df = df[df[col].astype(str) == val]
                break
    return df


# ── Network graph builder ─────────────────────────────────────────────────────
PATTERN_COLORS = {
    "FAN-OUT":        "#e67e22",
    "STRUCTURING":    "#f1c40f",
    "LAYERING":       "#9b59b6",
    "RAPID-MOVEMENT": "#e74c3c",
    "NORMAL":         "#95a5a6",
}

def build_network_graph(customer_id):
    """
    Build a Plotly figure showing the transaction network for a given customer.
    Returns (fig, feature_fig) — network graph and feature bar chart.
    Returns (None, None) if no transaction data available.
    """
    import plotly.graph_objects as go
    import networkx as nx
    import math

    if DF_TXN is None:
        return None, None

    all_txns = DF_TXN[DF_TXN["sender_customer_id"] == customer_id].copy()
    if all_txns.empty:
        return None, None

    # Only show suspicious (non-NORMAL) transactions in the graph
    txns = all_txns[all_txns["pattern_type"] != "NORMAL"].copy()
    if txns.empty:
        # Customer has no suspicious transactions — show a message
        empty = go.Figure(layout=go.Layout(
            paper_bgcolor="#1a1a1a", plot_bgcolor="#111",
            font=dict(color="#eee"),
            annotations=[dict(text="No suspicious transactions found for this customer.",
                              showarrow=False, font=dict(size=13, color="#aaa"),
                              xref="paper", yref="paper", x=0.5, y=0.5)],
            xaxis=dict(visible=False), yaxis=dict(visible=False),
        ))
        return empty, None

    # ── Build directed graph ──────────────────────────────────────────────────
    G = nx.DiGraph()
    sender_acct = str(txns["sender_account_id"].iloc[0])
    G.add_node(sender_acct, node_type="sender", label="Customer\n…" + sender_acct[-4:])

    for _, row in txns.iterrows():
        dst  = str(row["receiver_account_id"])
        bank = row["receiver_bank_name"]
        pat  = row["pattern_type"]
        amt  = row["amount"]
        if not G.has_node(dst):
            G.add_node(dst, node_type="receiver", label=bank + "\n…" + dst[-4:])
        txn_id = str(row.get("txn_id", ""))
        if G.has_edge(sender_acct, dst):
            G[sender_acct][dst]["weight"]  += amt
            G[sender_acct][dst]["count"]   += 1
            G[sender_acct][dst]["txn_ids"].append(txn_id)
        else:
            G.add_edge(sender_acct, dst, weight=amt, count=1, pattern=pat,
                       currency=row["currency"], bank=bank, txn_ids=[txn_id])

    # ── Layout ────────────────────────────────────────────────────────────────
    if len(G.nodes) <= 2:
        pos = nx.shell_layout(G)
    else:
        pos = nx.spring_layout(G, seed=42, k=2.5 / math.sqrt(len(G.nodes)))

    # ── Edge traces (one per pattern type for legend) ─────────────────────────
    edge_groups = {}
    for u, v, data in G.edges(data=True):
        pat = data.get("pattern", "NORMAL")
        if pat not in edge_groups:
            edge_groups[pat] = {"x": [], "y": [], "widths": [], "texts": []}
        x0, y0 = pos[u]
        x1, y1 = pos[v]
        edge_groups[pat]["x"] += [x0, x1, None]
        edge_groups[pat]["y"] += [y0, y1, None]
        txn_ids = data.get("txn_ids", [])
        ids_str = ", ".join(txn_ids) if txn_ids else "—"
        edge_groups[pat]["texts"].append(
            f"<b>{data['bank']}</b><br>"
            f"Pattern: {pat}<br>"
            f"Amount: ${data['weight']:,.0f}  |  Txns: {data['count']}<br>"
            f"Txn ID(s): {ids_str}"
        )

    edge_traces = []
    for pat, eg in edge_groups.items():
        color = PATTERN_COLORS.get(pat, "#aaa")
        edge_traces.append(go.Scatter(
            x=eg["x"], y=eg["y"],
            mode="lines",
            line=dict(width=2, color=color),
            name=pat,
            hoverinfo="none",
            legendgroup=pat,
        ))

    # ── Node trace ────────────────────────────────────────────────────────────
    node_x, node_y, node_text, node_hover, node_color, node_size = [], [], [], [], [], []
    for node in G.nodes():
        x, y = pos[node]
        ntype = G.nodes[node].get("node_type", "receiver")
        label = G.nodes[node].get("label", node[-6:])
        node_x.append(x)
        node_y.append(y)
        node_text.append(label)
        if ntype == "sender":
            node_color.append("#2980b9")
            node_size.append(22)
            node_hover.append(f"<b>Customer Account</b><br>…{str(node)[-6:]}")
        else:
            out_deg = G.out_degree(node)
            in_deg  = G.in_degree(node)
            node_color.append("#c0392b" if in_deg > 1 else "#27ae60")
            node_size.append(14 + min(in_deg * 3, 14))
            node_hover.append(f"<b>{label}</b><br>Account: …{str(node)[-6:]}")

    node_trace = go.Scatter(
        x=node_x, y=node_y,
        mode="markers+text",
        marker=dict(size=node_size, color=node_color,
                    line=dict(width=1, color="#fff")),
        text=node_text,
        textposition="top center",
        textfont=dict(size=8),
        hovertext=node_hover,
        hoverinfo="text",
        showlegend=False,
    )

    n_txns    = len(txns)           # suspicious only
    n_rcv     = txns["receiver_account_id"].nunique()
    total_amt = txns["amount"].sum()
    n_total   = len(all_txns)       # all transactions including normal
    patterns  = txns["pattern_type"].unique()
    pat_str   = ", ".join(patterns) if len(patterns) else "None detected"

    fig = go.Figure(
        data=edge_traces + [node_trace],
        layout=go.Layout(
            title=dict(
                text=f"Suspicious Transactions — {n_txns} flagged / {n_total} total | "
                     f"{n_rcv} receivers | ${total_amt:,.0f}<br>"
                     f"<sup>Patterns: {pat_str}  (normal transactions hidden)</sup>",
                font=dict(size=13),
            ),
            showlegend=True,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
            hovermode="closest",
            margin=dict(l=20, r=20, t=80, b=20),
            xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            plot_bgcolor="#111",
            paper_bgcolor="#1a1a1a",
            font=dict(color="#eee"),
        ),
    )

    # ── Feature bar chart ─────────────────────────────────────────────────────
    feat_fig = None
    if DF_NET is not None and customer_id in DF_NET.index:
        row     = DF_NET.loc[customer_id]
        pop_med = DF_NET.median()
        feat_cols = ["fan_out_score", "structuring_score", "txn_velocity",
                     "cross_bank_ratio", "out_degree", "unique_banks_sent_to"]
        feat_labels = {
            "fan_out_score":       "Fan-Out Score",
            "structuring_score":   "Structuring Score",
            "txn_velocity":        "Txn Velocity",
            "cross_bank_ratio":    "Cross-Border Ratio",
            "out_degree":          "Receiver Count",
            "unique_banks_sent_to":"Banks Used",
        }
        available = [c for c in feat_cols if c in row.index]
        cust_vals = [row[c] for c in available]
        med_vals  = [pop_med[c] for c in available]
        labels    = [feat_labels.get(c, c) for c in available]

        feat_fig = go.Figure()
        feat_fig.add_trace(go.Bar(
            name="This customer",
            x=labels, y=cust_vals,
            marker_color="#e74c3c",
        ))
        feat_fig.add_trace(go.Bar(
            name="Population median",
            x=labels, y=med_vals,
            marker_color="#3498db",
            opacity=0.6,
        ))
        feat_fig.update_layout(
            barmode="group",
            title=dict(text="Network Features vs. Population Median", font=dict(size=12)),
            margin=dict(l=20, r=20, t=40, b=60),
            legend=dict(orientation="h", y=1.12),
            plot_bgcolor="#111",
            paper_bgcolor="#1a1a1a",
            font=dict(color="#eee", size=10),
            xaxis=dict(tickangle=-30),
            yaxis=dict(gridcolor="#333"),
        )

    return fig, feat_fig


# ── Pre-populate cluster cache at startup ─────────────────────────────────────
# Runs clustering for All / Business / Individual so the first user request
# is served from cache with no delay.
_startup_cluster_cache = {}
if DF_SS is not None:
    try:
        print("Pre-computing cluster cache at startup...")
        import ds_data_prep, lambda_ds_performance
        _sc_fig, _sc_stats, _sc_df = lambda_ds_performance.perform_clustering(DF_SS, "All", 4)
        _startup_cluster_cache = {
            "df_clustered":  _sc_df,
            "df_enriched":   _enrich_cluster_df(_sc_df),
            "scatter_fig":   _sc_fig,
            "stats":         _sc_stats,
            "customer_type": "All",
        }
        print(f"Startup cluster cache ready (All): {len(_sc_df):,} customers, 4 clusters")
    except Exception as _e:
        print(f"Startup clustering failed (non-fatal): {_e}")

_startup_cluster_cache_biz = {}
_startup_cluster_cache_ind = {}
if DF_SS is not None:
    try:
        import lambda_ds_performance as _ldp
        _biz_fig, _biz_stats, _biz_df = _ldp.perform_clustering(DF_SS, "Business", 4)
        _startup_cluster_cache_biz = {
            "df_clustered":  _biz_df,
            "df_enriched":   _enrich_cluster_df(_biz_df),
            "scatter_fig":   _biz_fig,
            "stats":         _biz_stats,
            "customer_type": "Business",
        }
        print(f"Startup cluster cache ready (Business): {len(_biz_df):,} customers, 4 clusters")
        _ind_fig, _ind_stats, _ind_df = _ldp.perform_clustering(DF_SS, "Individual", 4)
        _startup_cluster_cache_ind = {
            "df_clustered":  _ind_df,
            "df_enriched":   _enrich_cluster_df(_ind_df),
            "scatter_fig":   _ind_fig,
            "stats":         _ind_stats,
            "customer_type": "Individual",
        }
        print(f"Startup cluster cache ready (Individual): {len(_ind_df):,} customers, 4 clusters")
    except Exception as _e:
        print(f"Startup Business/Individual clustering failed (non-fatal): {_e}")

# ── Build DF_CLUSTER_LABELS from startup cache when CSV labels don't match ─────
# If the static CSV has no overlap with the current rule sweep data (e.g. synth),
# derive cluster labels from the All-customers startup clustering instead.
if DF_CLUSTER_LABELS is None or (
    DF_RULE_SWEEP is not None
    and "customer_id" in (DF_CLUSTER_LABELS.columns if DF_CLUSTER_LABELS is not None else [])
    and len(set(DF_CLUSTER_LABELS["customer_id"]) & set(DF_RULE_SWEEP["customer_id"])) == 0
):
    _sc_df = _startup_cluster_cache.get("df_clustered")
    if _sc_df is not None and "customer_id" in _sc_df.columns:
        # clusters are 0-based from KMeans; store 1-based to match _filter_by_cluster expectation
        DF_CLUSTER_LABELS = _sc_df[["customer_id", "cluster"]].copy()
        DF_CLUSTER_LABELS["cluster"] = DF_CLUSTER_LABELS["cluster"] + 1
        print(f"Cluster labels: derived from startup cache ({len(DF_CLUSTER_LABELS):,} customers)")

COL_MAP = {
    "AVG_TRXNS_WEEK":   "avg_num_trxns",
    "AVG_TRXN_AMT":     "avg_trxn_amt",
    "TRXN_AMT_MONTHLY": "trxn_amt_monthly",
}

# SAR simulation uses ds_segmentation_data column names (slightly different from main DF)
SAR_COL_MAP = {
    "AVG_TRXNS_WEEK":   "avg_num_trxns",
    "AVG_TRXN_AMT":     "avg_weekly_trxn_amt",
    "TRXN_AMT_MONTHLY": "trxn_amt_monthly",
}

orchestrator = OrchestratorAgent()

# Pin model in GPU — overrides Ollama's 5-minute idle eviction for the session.
try:
    import urllib.request as _urlreq, json as _json
    _pin_url  = OLLAMA_BASE_URL.replace("/v1", "") + "/api/generate"
    _pin_data = _json.dumps({"model": OLLAMA_MODEL, "keep_alive": -1}).encode()
    _urlreq.urlopen(_urlreq.Request(_pin_url, data=_pin_data,
                                    headers={"Content-Type": "application/json"}), timeout=10)
    print(f"Ollama: model '{OLLAMA_MODEL}' pinned in GPU (keep_alive=-1)")
except Exception as _e:
    print(f"Ollama: could not pin model — {_e}")

# ── Suggested prompts ─────────────────────────────────────────────────────────
DEMO_PROMPTS = [
    "What are sanctions lists",
    "What is threshold tuning",
    "What is dynamic segmentation",
    "What is OFAC",
    "What are the main differences between US and Canada in AML regulations",
    "What is SAR backtesting",
    "How does ARIA help with SAR backtesting",
]

SUGGESTED_PROMPTS = [
    # Threshold Tuning
    "Show FP/FN trade-off for Business customers by monthly transaction amount",
    "Run SAR backtest for Individual customers using average weekly transactions",
    "What is the crossover threshold for Business customers on average transaction amount?",
    # Dynamic Segmentation
    "Cluster Business customers by transaction behavior",
    "Show me the behavioral segments for Individual customers",
    "Which cluster of Business customers has the highest transaction volume?",
    # Rule-Level Sweep
    "Show SAR backtest for Activity Deviation ACH rule",
    "Run a 2D sweep for Elder Abuse varying floor amount and age threshold",
    "Show Elder Abuse SAR backtest for Cluster 4",
]

# ── Welcome message ───────────────────────────────────────────────────────────
_WELCOME = (
    f"Hello! I'm **ARIA** — Agentic Risk Intelligence for AML — powered by **{OLLAMA_MODEL}**.\n\n"
    "I can help you with:\n"
    "- **Threshold tuning** — analyze how FP/FN trade-offs shift as alert thresholds change\n"
    "- **Dynamic Segmentation** — cluster customers into behavioral segments using K-Means\n"
    "- **AML policy Q&A** — answer compliance questions from the knowledge base\n\n"
    f"Dataset loaded: **{_total:,} accounts** ({_biz_count:,} Business / {_ind_count:,} Individual) "
    f"| **{_alert_count:,} alerts** | **{_fp_count:,} false positives**\n\n"
    "**Tip:** Use the sidebar buttons on the left to open the **SAR Priority Worklist** "
    "(ranked alerts with network graph viewer) and **Segment Customer Drilldown**. "
    "Click a suggested prompt below or type your question."
)

INITIAL_MESSAGES = [{"role": "assistant", "content": _WELCOME}]


# ── Tool helpers ──────────────────────────────────────────────────────────────
def compute_threshold_stats(df_seg, threshold):
    """
    Compute threshold sweep and return:
      - a pre-written factual interpretation (Python-generated, always accurate)
      - the raw sweep table appended for reference

    The model is instructed to copy the PRE-COMPUTED ANALYSIS verbatim and only
    add one AML insight sentence — this eliminates hallucinated numbers.
    """
    import math as _math
    cur_val = df_seg[threshold].median()   # use median as operational center
    raw_step = cur_val / 4
    mag = 10 ** _math.floor(_math.log10(max(raw_step, 1)))
    n = raw_step / mag
    if n < 1.5:   nice = 1
    elif n < 3.5: nice = 2
    elif n < 7.5: nice = 5
    else:         nice = 10
    step  = int(nice * mag)
    t_min = max(step, int(cur_val - 4 * step))
    t_max = int(cur_val + 4 * step)

    sweep = []
    t = t_min
    while t <= t_max + step:
        t_r = round(t, 2)
        fp = df_seg[(df_seg[threshold] >= t_r) & (df_seg["false_positives"] == 1)].shape[0]
        fn = df_seg[(df_seg[threshold] <  t_r) & (df_seg["false_negatives"] == 1)].shape[0]
        sweep.append((t_r, fp, fn))
        t += step

    # ── Derive key facts ──────────────────────────────────────────────────────
    max_fp = sweep[0][1]
    t_last = sweep[-1][0]   # actual last threshold in sweep (may be t_max + step)
    max_fn = sweep[-1][2]

    fn_first_nonzero = next(((t, fn) for t, fp, fn in sweep if fn > 0), None)
    fp_first_zero    = next(((t, fp) for t, fp, fn in sweep if fp == 0), None)
    fn_zero_end      = max(t_min, round(fn_first_nonzero[0] - step, 2)) if fn_first_nonzero else t_max

    crossover = min(sweep, key=lambda x: abs(x[1] - x[2]))

    optimal = [(t, fp, fn) for t, fp, fn in sweep
               if fp <= max_fp * 0.20 and fn <= max_fn * 0.20]

    # ── Build pre-written factual interpretation ──────────────────────────────
    lines = ["=== PRE-COMPUTED ANALYSIS (copy this verbatim, do not alter numbers) ==="]

    lines.append(f"### Threshold Tuning — False Positive / False Negative Trade-off\n")

    lines.append(f"**False Positives (FP)**")
    lines.append(f"- At the lowest threshold ({t_min}): **{max_fp} FPs**")
    if fp_first_zero:
        lines.append(f"- FPs reach zero at threshold **{fp_first_zero[0]}**")
    else:
        lines.append(f"- FPs do not reach zero within the sweep range ({t_min}–{t_max})")
    lines.append("")

    lines.append(f"**False Negatives (FN)**")
    if fn_first_nonzero:
        if fn_zero_end > t_min:
            lines.append(f"- FNs are zero from threshold {t_min} up to **{fn_zero_end}**")
        else:
            lines.append(f"- FNs are non-zero even at the lowest threshold ({t_min}) — some customers fall below the sweep floor")
        lines.append(f"- FNs first appear at threshold **{fn_first_nonzero[0]}** (FN={fn_first_nonzero[1]})")
        lines.append(f"- FNs reach **{max_fn}** at the highest threshold ({t_last})")
    else:
        lines.append(f"- FNs remain zero across the entire sweep range")
    lines.append("")

    lines.append(f"**Crossover Point** — threshold **{crossover[0]}** (FP={crossover[1]}, FN={crossover[2]})")
    lines.append("")

    if optimal:
        lines.append(
            f"**Optimal Zone** (FP and FN both below 20% of max): threshold **{optimal[0][0]}** to **{optimal[-1][0]}**"
        )
    else:
        lines.append("**Optimal Zone**: no single threshold achieves both FP and FN below 20% of their maximums simultaneously.")

    lines.append("\n*(Detailed sweep chart shown below.)*")
    lines.append("=== END PRE-COMPUTED ANALYSIS ===")

    return "\n".join(lines)


def compute_segment_stats(df):
    total_alerts = int(df["alerts"].sum())
    total_accounts = len(df)
    lines = ["=== PRE-COMPUTED SEGMENT STATS (copy verbatim, do not compute new numbers) ==="]
    lines.append("### Segment Overview\n")
    for seg_id, name in [(0, "Business"), (1, "Individual")]:
        seg = df[df["dynamic_segment"] == seg_id]
        n         = len(seg)
        alerts    = int(seg["alerts"].sum())
        fp        = int(seg["false_positives"].sum())
        fn        = int(seg["false_negatives"].sum())
        tp        = alerts - fp
        fp_rate   = round(100 * fp / alerts, 1) if alerts > 0 else 0
        sar_rate  = round(100 * tp / alerts, 1) if alerts > 0 else 0
        acct_pct  = round(100 * n / total_accounts, 1) if total_accounts > 0 else 0
        alert_pct = round(100 * alerts / total_alerts, 1) if total_alerts > 0 else 0
        lines.append(f"**{name}**")
        lines.append(f"- Accounts: **{n:,}** ({acct_pct}% of total)")
        lines.append(f"- Alerts: **{alerts:,}** ({alert_pct}% of all alerts)")
        lines.append(f"- SARs (True Positives): **{tp:,}** (SAR rate={sar_rate}% of alerts)")
        lines.append(f"- False Positives: **{fp:,}** (FP rate={fp_rate}% of alerts)")
        lines.append(f"- False Negatives: **{fn:,}**")
        lines.append("")
    lines.append("=== END PRE-COMPUTED SEGMENT STATS ===")
    return "\n".join(lines)


def compute_sar_backtest(df_sar_seg, sar_col, segment_name):
    """
    Sweep thresholds across SAR customers and report how many simulated SARs
    would be caught vs. missed at each threshold level.

    Returns PRE-COMPUTED text for the model to copy verbatim.
    """
    df = df_sar_seg.dropna(subset=[sar_col]).copy()
    total_sars = int(df["is_sar"].sum())
    total_alerted = len(df)

    if total_sars == 0:
        return "No simulated SARs found for this segment and threshold column."

    t_min = df[sar_col].min()
    t_max = df[sar_col].max()
    step  = max(1, int((t_max - t_min) / 100))

    sweep = []
    t = t_min
    while t <= t_max + step:
        caught = int(((df[sar_col] >= t) & (df["is_sar"] == 1)).sum())
        missed = total_sars - caught
        sweep.append((round(t, 2), caught, missed))
        t += step

    # ── Key statistics ────────────────────────────────────────────────────────
    # Threshold where SAR catch rate first drops below 90% / 80% / 50%
    def threshold_for_rate(target_rate):
        target = int(total_sars * target_rate)
        hit = next(((t, c, m) for t, c, m in sweep if c <= target), None)
        return hit

    t90 = threshold_for_rate(0.90)   # threshold where we drop below 90% catch
    t80 = threshold_for_rate(0.80)
    t50 = threshold_for_rate(0.50)

    # Threshold where SARs first start being missed
    first_miss = next(((t, c, m) for t, c, m in sweep if m > 0), None)

    lines = ["=== PRE-COMPUTED SAR BACKTEST (copy this verbatim, do not alter numbers) ==="]

    lines.append(f"### SAR Backtest — {segment_name} / {sar_col}\n")
    lines.append(f"**Population:** {total_alerted:,} alerted customers | **SARs:** {total_sars:,} ({round(100*total_sars/total_alerted,1)}% SAR filing rate)\n")

    lines.append(f"**Sweep Results**")
    lines.append(f"- At lowest threshold ({sweep[0][0]}): **{sweep[0][1]} SARs caught** (100%), 0 missed")
    if first_miss:
        lines.append(f"- SARs first missed at threshold **{first_miss[0]}** ({first_miss[2]} missed)")

    if t90:
        prev = next((s for s in reversed(sweep) if s[0] < t90[0] and s[1] > int(total_sars*0.90)), None)
        t90_keep = prev[0] if prev else sweep[0][0]
        lines.append(f"- To keep ≥90% SAR catch rate: threshold ≤ **{t90_keep}** ({int(total_sars*0.90)+1} of {total_sars} caught)")
    if t80:
        prev = next((s for s in reversed(sweep) if s[0] < t80[0] and s[1] > int(total_sars*0.80)), None)
        t80_keep = prev[0] if prev else sweep[0][0]
        lines.append(f"- To keep ≥80% SAR catch rate: threshold ≤ **{t80_keep}** ({int(total_sars*0.80)+1} of {total_sars} caught)")
    if t50:
        prev = next((s for s in reversed(sweep) if s[0] < t50[0] and s[1] > int(total_sars*0.50)), None)
        t50_keep = prev[0] if prev else sweep[0][0]
        lines.append(f"- To keep ≥50% SAR catch rate: threshold ≤ **{t50_keep}** ({int(total_sars*0.50)+1} of {total_sars} caught)")

    lines.append(f"- At highest threshold ({sweep[-1][0]}): **{sweep[-1][1]} caught**, {sweep[-1][2]} missed")
    lines.append("\n*(Detailed sweep chart shown below.)*")
    lines.append("=== END PRE-COMPUTED SAR BACKTEST ===")

    return "\n".join(lines)


# ── Cluster filter helper ─────────────────────────────────────────────────────
def _filter_by_cluster(df_rule_sweep, cluster):
    """
    Return df_rule_sweep filtered to customers in a specific behavioral cluster.
    cluster: int (1-based) or None (no filtering).
    Uses static DF_CLUSTER_LABELS (1-indexed, built from alerted population).
    """
    if cluster is None or DF_CLUSTER_LABELS is None:
        return df_rule_sweep
    cluster = int(cluster)
    ids      = DF_CLUSTER_LABELS[DF_CLUSTER_LABELS["cluster"] == cluster]["customer_id"]
    filtered = df_rule_sweep[df_rule_sweep["customer_id"].isin(ids)]
    print(f"[cluster filter] cluster={cluster}: {len(filtered)} / {len(df_rule_sweep)} rows")
    return filtered


# ── Tool executor ─────────────────────────────────────────────────────────────
_cluster_cache       = _startup_cluster_cache  # pre-populated at startup; updated on each cluster run
_thread_local        = threading.local()        # per-request: .current_query, .last_2d_state
_last_cluster_result    = ""   # last clustering model response (for multi-turn follow-ups)
_last_cluster_raw_stats = ""   # raw pre-computed stats block — more reliable for comparison queries

def _json_safe(obj):
    """Recursively convert numpy types to JSON-serializable Python types."""
    import numpy as np
    if isinstance(obj, dict):
        return {k: _json_safe(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_json_safe(i) for i in obj]
    if isinstance(obj, np.integer):
        return int(obj)
    if isinstance(obj, np.floating):
        return float(obj)
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    return obj

def tool_executor(tool_name, tool_input):
    """Execute a tool called by an agent. Returns (result_text, fig_or_None)."""
    global DF_SS, _cluster_cache

    if tool_name == "threshold_tuning":
        segment = tool_input.get("segment", "Business")
        raw_col = tool_input.get("threshold_column", "AVG_TRXNS_WEEK")
        col     = COL_MAP.get(raw_col)
        if col is None:
            return (f"'{raw_col}' is not a valid threshold_column. "
                    f"Valid options: AVG_TRXNS_WEEK, AVG_TRXN_AMT, TRXN_AMT_MONTHLY."), None
        df_seg  = DF_BUSINESS if segment == "Business" else DF_INDIVIDUAL
        stats   = compute_threshold_stats(df_seg, col)
        line_fig, _ = lambda_ds_performance.plot_thresholds_tuning(df_seg, col, 0.1, segment)
        tbl_fig = make_figures.threshold_tuning_figure(df_seg, col, segment)
        return stats, (line_fig, tbl_fig)

    elif tool_name == "segment_stats":
        tbl_fig = make_figures.segment_stats_figure(DF)
        return compute_segment_stats(DF), tbl_fig

    elif tool_name == "sar_backtest":
        if DF_SAR is None:
            return "SAR simulation data not found. Run python simulate_sars.py first.", None
        segment    = tool_input.get("segment", "Business")
        raw_col    = tool_input.get("threshold_column", "TRXN_AMT_MONTHLY")
        col        = SAR_COL_MAP.get(raw_col)
        if col is None:
            return (f"'{raw_col}' is not a valid threshold_column. "
                    f"Valid options: AVG_TRXNS_WEEK, AVG_TRXN_AMT, TRXN_AMT_MONTHLY."), None
        df_sar_seg = DF_SAR_BUSINESS if segment == "Business" else DF_SAR_INDIVIDUAL
        stats      = compute_sar_backtest(df_sar_seg, col, segment)
        tbl_fig    = make_figures.sar_backtest_figure(df_sar_seg, col, segment)
        return stats, tbl_fig

    elif tool_name == "rule_2d_sweep":
        if DF_RULE_SWEEP is None:
            return "Rule sweep data not found. Run python prepare_rule_sweep_data.py first.", None
        risk_factor = tool_input.get("risk_factor", "")
        param1      = tool_input.get("sweep_param_1") or None
        param2      = tool_input.get("sweep_param_2") or None
        cluster     = tool_input.get("cluster", None)
        df_sweep    = _filter_by_cluster(DF_RULE_SWEEP, cluster)
        text, grid  = lambda_rule_analysis.compute_rule_2d_sweep(df_sweep, risk_factor, param1, param2)
        heatmap     = make_figures.rule_2d_heatmap(grid) if grid else None
        ranked_tbl  = make_figures.rule_2d_ranked_table(grid) if grid else None
        if grid:
            if not hasattr(_thread_local, 'last_2d_state'):
                _thread_local.last_2d_state = {}
            _thread_local.last_2d_state.update({
                "grid":        grid,
                "risk_factor": risk_factor,
                "param1":      grid["param1"],
                "param2":      grid["param2"],
                "cluster":     cluster,
                "ts":          time.time(),
            })
        # If cluster-filtered, also show rule alert distribution by cluster
        if cluster is not None and heatmap is not None and DF_CLUSTER_LABELS is not None:
            dist_fig = make_figures.rule_alerts_by_cluster(
                DF_RULE_SWEEP, DF_CLUSTER_LABELS, grid["rf_name"] if grid else risk_factor, cluster
            )
            if dist_fig is not None:
                figs = [f for f in [ranked_tbl, dist_fig] if f is not None]
                return text, tuple(figs) if len(figs) > 1 else (figs[0] if figs else None)
        # Heatmap lives in the side panel -- only show ranked table in chat
        return text, ranked_tbl

    elif tool_name == "cluster_rule_summary":
        if DF_RULE_SWEEP is None:
            return "Rule sweep data not found. Run python prepare_rule_sweep_data.py first.", None
        cluster = tool_input.get("cluster")
        if cluster is None:
            return "cluster parameter is required (integer 1–4).", None
        df_filtered = _filter_by_cluster(DF_RULE_SWEEP, cluster)
        stats   = lambda_rule_analysis.cluster_rule_summary_text(df_filtered, cluster)
        tbl_fig = make_figures.rule_list_figure(df_filtered)
        return stats, tbl_fig

    elif tool_name == "list_rules":
        if DF_RULE_SWEEP is None:
            return "Rule sweep data not found. Run python prepare_rule_sweep_data.py first.", None
        tbl_fig = make_figures.rule_list_figure(DF_RULE_SWEEP)
        return lambda_rule_analysis.list_rules_text(DF_RULE_SWEEP), tbl_fig

    elif tool_name == "rule_sar_backtest":
        if DF_RULE_SWEEP is None:
            return "Rule sweep data not found. Run python prepare_rule_sweep_data.py first.", None
        risk_factor = tool_input.get("risk_factor", tool_input.get("rule_code", ""))
        sweep_param = tool_input.get("sweep_param", None)
        cluster     = tool_input.get("cluster", None)
        df_sweep    = _filter_by_cluster(DF_RULE_SWEEP, cluster)
        stats   = lambda_rule_analysis.compute_rule_sar_sweep(df_sweep, risk_factor, sweep_param)
        tbl_fig = make_figures.rule_sweep_figure(df_sweep, risk_factor, sweep_param)
        return stats, tbl_fig

    elif tool_name == "cluster_threshold_analysis":
        if DF_SS is None or DF_SAR is None:
            return "Segmentation or SAR data not available.", None
        import lambda_cluster_threshold
        segment     = tool_input.get("segment", "Business")
        raw_col     = tool_input.get("threshold_column", "AVG_TRXNS_WEEK")
        n_clusters  = int(tool_input.get("n_clusters", 4))
        target_rate = float(tool_input.get("target_sar_rate", 0.90))
        text, fig = lambda_cluster_threshold.cluster_threshold_analysis(
            DF_SS, DF_SAR, segment, raw_col, n_clusters, target_rate
        )
        return text, fig

    elif tool_name == "alerts_distribution":
        bar_fig = lambda_ds_performance.alerts_distribution(DF)
        tbl_fig = make_figures.segment_stats_figure(DF)
        return compute_segment_stats(DF), (bar_fig, tbl_fig)

    elif tool_name == "cluster_analysis":
        customer_type = tool_input.get("customer_type", "All")
        n_clusters    = max(2, min(6, int(tool_input.get("n_clusters", 4))))
        df_src        = DF_SS if DF_SS is not None else DF
        scatter_fig, stats, df_clustered = lambda_ds_performance.perform_clustering(
            df_src, customer_type, n_clusters
        )
        _cluster_cache.update({
            "df_clustered":  df_clustered,
            "df_enriched":   _enrich_cluster_df(df_clustered),
            "scatter_fig":   scatter_fig,
            "stats":         stats,
            "customer_type": customer_type,
        })
        global _last_cluster_raw_stats
        _last_cluster_raw_stats = stats
        treemap_fig = lambda_ds_performance.smartseg_tree_dynamic(
            df_clustered, customer_type, dims=_DS_DIMS, df_rule_sweep=DF_RULE_SWEEP
        )
        return stats, (scatter_fig, treemap_fig)

    elif tool_name == "prepare_segmentation_data":
        import ds_data_prep
        try:
            df_out = ds_data_prep.prepare_data()
            DF_SS  = df_out
            summary = (
                f"Data prep complete — {ds_data_prep.OUTPUT_SYNTH if ds_data_prep._USE_SYNTH else ds_data_prep.OUTPUT_FILE}\n"
                f"Rows: {len(df_out):,} | Columns: {len(df_out.columns)}\n"
                f"customer_type: {df_out['customer_type'].value_counts().to_dict()}"
            )
            return summary, None
        except Exception as e:
            return f"Data prep error: {e}", None

    elif tool_name == "ds_cluster_analysis":
        import ds_data_prep
        if DF_SS is None:
            print("ds_cluster_analysis: DF_SS not loaded — running prepare_data() first ...")
            DF_SS = ds_data_prep.prepare_data()
        customer_type   = tool_input.get("customer_type", "All")
        n_clusters      = max(2, min(6, int(tool_input.get("n_clusters", 4))))
        filter_clusters = tool_input.get("filter_clusters", None)

        # If the user didn't mention a specific cluster count, ignore the model's choice and use 4.
        _cq = getattr(_thread_local, 'current_query', '')
        _n_match = re.search(r'\b([2-8])\s*clusters?\b', _cq, re.IGNORECASE)
        if not _n_match:
            _n_match = re.search(r'\binto\s+([2-8])\b', _cq, re.IGNORECASE)
        if not _n_match:
            _n_match = re.search(r'\bonly\s+([2-8])\b', _cq, re.IGNORECASE)
        if _n_match:
            n_clusters = max(2, min(6, int(_n_match.group(1))))
        else:
            n_clusters = 4  # user didn't specify — always default to 4

        if filter_clusters:
            # Second call: use startup prebuilt for this customer_type — never the shared
            # _cluster_cache which may hold a different user's clustering run.
            _prebuilt_for_filter = {
                "All":        _startup_cluster_cache,
                "Business":   _startup_cluster_cache_biz,
                "Individual": _startup_cluster_cache_ind,
            }.get(customer_type, _startup_cluster_cache)
            if not _prebuilt_for_filter:
                return "No clustering data available for filtering. Run clustering first.", None
            df_clustered = _prebuilt_for_filter["df_clustered"]
            scatter_fig  = _prebuilt_for_filter["scatter_fig"]
            stats        = _prebuilt_for_filter["stats"]

            # Filter scatter: keep only traces for requested clusters
            import plotly.graph_objects as go
            keep_names = {f"Cluster {c}" for c in filter_clusters}
            filtered_scatter = go.Figure(
                data=[t for t in scatter_fig.data if t.name in keep_names],
                layout=scatter_fig.layout,
            )
            filtered_scatter.update_layout(title=filtered_scatter.layout.title.text +
                                           f" [filtered: clusters {filter_clusters}]")

            # Filter treemap: keep only rows for requested clusters (0-based internally)
            cluster_vals = sorted(df_clustered["cluster"].unique())
            keep_0based  = {cluster_vals[c - 1] for c in filter_clusters
                            if 1 <= c <= len(cluster_vals)}
            df_filtered  = df_clustered[df_clustered["cluster"].isin(keep_0based)].copy()
            treemap_fig  = lambda_ds_performance.smartseg_tree_dynamic(
                df_filtered, f"{customer_type} (clusters {filter_clusters})", dims=_DS_DIMS, df_rule_sweep=DF_RULE_SWEEP
            )
            return f"Filtered to clusters {filter_clusters}.\n\n{stats}", (filtered_scatter, treemap_fig)

        else:
            # Use startup cache if available and n_clusters matches (avoids recomputing)
            _prebuilt = {
                "All":        _startup_cluster_cache,
                "Business":   _startup_cluster_cache_biz,
                "Individual": _startup_cluster_cache_ind,
            }.get(customer_type, {})
            if _prebuilt and n_clusters == 4:
                scatter_fig  = _prebuilt["scatter_fig"]
                stats        = _prebuilt["stats"]
                df_clustered = _prebuilt["df_clustered"]
                _cluster_cache.update(_prebuilt)
                treemap_fig  = lambda_ds_performance.smartseg_tree_dynamic(
                    df_clustered, customer_type, dims=_DS_DIMS, df_rule_sweep=DF_RULE_SWEEP
                )
                stats_table = make_figures.cluster_stats_table(df_clustered, customer_type)
                figs = (stats_table, scatter_fig, treemap_fig) if stats_table is not None else (scatter_fig, treemap_fig)
                _last_cluster_raw_stats = stats
                return stats, figs

            # Cache miss or non-default n_clusters: run full clustering
            scatter_fig, stats, df_clustered = lambda_ds_performance.perform_clustering(
                DF_SS, customer_type, n_clusters
            )
            _cluster_cache.update({
                "df_clustered":  df_clustered,
                "df_enriched":   _enrich_cluster_df(df_clustered),
                "scatter_fig":   scatter_fig,
                "stats":         stats,
                "customer_type": customer_type,
            })
            treemap_fig  = lambda_ds_performance.smartseg_tree_dynamic(
                df_clustered, customer_type, dims=_DS_DIMS, df_rule_sweep=DF_RULE_SWEEP
            )
            stats_table  = make_figures.cluster_stats_table(df_clustered, customer_type)
            figs = (stats_table, scatter_fig, treemap_fig) if stats_table is not None else (scatter_fig, treemap_fig)
            _last_cluster_raw_stats = stats
            return stats, figs

    elif tool_name == "ofac_screening":
        filter_type = tool_input.get("filter_type", "all")
        return lambda_ofac.ofac_screening(filter_type=filter_type)

    elif tool_name == "ofac_name_lookup":
        name      = tool_input.get("name", "")
        threshold = float(tool_input.get("threshold", 85))
        return lambda_ofac.screen_name(name, threshold=threshold)

    return f"Unknown tool: {tool_name}", None


# ── Chat export ──────────────────────────────────────────────────────────────

def _add_formatted_runs(paragraph, text):
    """Split text on **bold** markers and add styled runs to a docx paragraph."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _messages_to_docx(messages) -> bytes:
    """Convert chat message list to a Word (.docx) document and return raw bytes."""
    import docx as _docx
    from docx.shared import Pt, RGBColor

    doc = _docx.Document()
    doc.core_properties.title = "ARIA Analysis Session"

    title = doc.add_heading("ARIA Analysis Session", level=0)
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    sub = doc.add_paragraph(f"Exported: {ts}")
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for msg in messages:
        role    = msg.get("role", "")
        content = msg.get("content", "")
        if isinstance(content, list):
            content = " ".join(
                b.get("text", "") for b in content
                if isinstance(b, dict) and b.get("type") == "text"
            )
        if not content:
            continue

        # Role header
        label_para = doc.add_paragraph()
        label_run  = label_para.add_run("You" if role == "user" else "ARIA")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = (
            RGBColor(0x1a, 0x73, 0xe8) if role == "user" else RGBColor(0x1e, 0x88, 0x5a)
        )

        # Content — line by line with basic markdown handling
        for line in content.split("\n"):
            s = line.rstrip()
            if s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_runs(p, s[2:])
            elif s in ("---", "***", "___"):
                doc.add_paragraph()
            else:
                p = doc.add_paragraph()
                _add_formatted_runs(p, s)

        doc.add_paragraph()  # spacer between turns

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Message compaction ───────────────────────────────────────────────────────
_MAX_CHAT_MESSAGES = 30

def _compact_chart_figures(messages, keep_last_n=1):
    """
    For all but the last `keep_last_n` chart messages, replace the Plotly figure
    with an empty placeholder.  Both trace data AND layout are stripped so that
    old chart messages add near-zero bytes to the messages payload.

    Keeping layout was the original design, but even an empty-trace heatmap
    layout can be 100-500 KB; across many turns this causes React reconciliation
    lag and GC pauses while the user is typing.

    The container style (height/width) is preserved so the chat layout stays stable.
    Only affects {"type": "graph", "props": {"figure": {...}}} blocks.
    """
    chart_indices = [
        i for i, m in enumerate(messages)
        if isinstance(m.get("content"), list)
        and any(b.get("type") == "graph" and "props" in b for b in m["content"])
    ]
    trim_indices = set(chart_indices[:-keep_last_n]) if len(chart_indices) > keep_last_n else set()
    if not trim_indices:
        return messages

    result = list(messages)
    for i in trim_indices:
        msg = dict(result[i])
        new_content = []
        for block in msg["content"]:
            if block.get("type") == "graph" and "props" in block:
                props = dict(block["props"])
                props["figure"] = {"data": [], "layout": {}}
                block = dict(block)
                block["props"] = props
            new_content.append(block)
        msg["content"] = new_content
        result[i] = msg
    return result


# ── Chart content builder ─────────────────────────────────────────────────────
def _chart_content(tool_name, tool_input, fig):
    """Convert a tool result into (message_blocks, chart_list).

    message_blocks: text-only content for the messages prop — NO Plotly JSON.
                    Keeping figure dicts out of messages eliminates React
                    reconciliation lag on every keystroke.
    chart_list:     [{"title": str, "figure": dict, "height": int}] passed to
                    charts-panel-store and rendered by render_charts_panel().
    """
    blocks = []
    chart_list = []

    if tool_name == "threshold_tuning":
        seg = tool_input.get("segment", "")
        col = tool_input.get("threshold_column", "")
        figs   = list(fig) if isinstance(fig, tuple) else [fig]
        labels = [f"Threshold Tuning — {seg} / {col}", f"Threshold Table — {seg} / {col}"][:len(figs)]

    elif tool_name == "alerts_distribution":
        figs   = list(fig) if isinstance(fig, tuple) else [fig]
        labels = ["Alerts & False Positives by Segment", "Segment Statistics Table"][:len(figs)]

    elif tool_name == "segment_stats":
        figs   = [fig]
        labels = ["Segment Statistics"]

    elif tool_name == "sar_backtest":
        seg = tool_input.get("segment", "")
        col = tool_input.get("threshold_column", "")
        figs   = [fig]
        labels = [f"SAR Backtest — {seg} / {col}"]

    elif tool_name == "list_rules":
        figs   = [fig]
        labels = ["AML Rule Performance Overview"]

    elif tool_name == "rule_2d_sweep":
        from lambda_rule_analysis import RULE_CATALOGUE, _match_rule
        rf      = tool_input.get("risk_factor", "")
        p1      = tool_input.get("sweep_param_1") or None
        p2      = tool_input.get("sweep_param_2") or None
        cluster = tool_input.get("cluster", None)
        _, entry = _match_rule(rf)
        if entry and (p1 is None or p2 is None):
            d1, d2 = entry.get("default_2d", (None, None))
            if p1 is None: p1 = d1
            if p2 is None: p2 = d2
        cluster_tag = f" [Cluster {cluster}]" if cluster else ""
        if isinstance(fig, tuple):
            figs = list(fig)
            base_labels = [
                f"2D Sweep Heatmap — {rf}{cluster_tag}",
                f"2D Sweep Ranked Table — {rf}{cluster_tag}",
            ]
            if len(figs) > 2:
                base_labels.append(f"{rf} — Alerts by Cluster")
            labels = base_labels[:len(figs)]
        else:
            figs   = [fig]
            labels = [f"2D Sweep — {rf} ({p1 or '?'} x {p2 or '?'}){cluster_tag}"]

    elif tool_name == "rule_sar_backtest":
        from lambda_rule_analysis import RULE_CATALOGUE, _match_rule
        rf      = tool_input.get("risk_factor", tool_input.get("rule_code", ""))
        sp      = tool_input.get("sweep_param") or None
        cluster = tool_input.get("cluster", None)
        _, entry = _match_rule(rf)
        if sp is None and entry:
            sp = entry["default_sweep"]
        cluster_tag = f" [Cluster {cluster}]" if cluster else ""
        figs   = [fig]
        labels = [f"Rule SAR Sweep — {rf} / {sp or 'default'}{cluster_tag}"]

    elif tool_name == "ofac_screening":
        title = "OFAC Sanctions Screening"
        if fig is not None:
            fig_h = fig.layout.height if fig.layout.height else 400
            chart_list.append({"title": title, "figure": fig.to_dict(), "height": fig_h})
        blocks.append({"type": "text", "text": f"**{title}** — chart shown below."})
        return blocks, chart_list

    elif tool_name in ("cluster_analysis", "ds_cluster_analysis"):
        ct = tool_input.get("customer_type", "All")
        if isinstance(fig, tuple) and len(fig) == 3:
            figs   = [fig[0]]   # stats_table → charts panel
            labels = [f"Cluster Summary — {ct}"]
        elif isinstance(fig, tuple):
            figs   = []
            labels = []
        else:
            figs   = [fig]
            labels = [f"Cluster Analysis — {ct}"]
        blocks.append({"type": "text", "text": "📊 Use **Segment Customer Drilldown** and **Cluster Scatter Plot** in the left sidebar to explore the segments interactively."})

    elif tool_name == "cluster_threshold_analysis":
        seg = tool_input.get("segment", "")
        col = tool_input.get("threshold_column", "")
        figs   = [fig]
        labels = [f"Adaptive Thresholds — {seg} / {col}"]

    else:
        figs   = [fig] if not isinstance(fig, tuple) else list(fig)
        labels = [tool_name] * len(figs)

    for label, f in zip(labels, figs):
        if f is None:
            continue
        fig_h = f.layout.height if f.layout.height else 460
        chart_list.append({"title": label, "figure": f.to_dict(), "height": fig_h})
    # Do not inject "See chart/table below." — chart renders in UI; injecting pollutes model history.

    return blocks, chart_list


# ── Dash app ──────────────────────────────────────────────────────────────────
server = flask.Flask(__name__)
app = Dash(
    __name__,
    server=server,
    routes_pathname_prefix="/",
    external_stylesheets=[dbc.themes.BOOTSTRAP],
    title="ARIA",
)

# ── Layout ────────────────────────────────────────────────────────────────────
_sidebar = dbc.Card([
    dbc.CardBody([
        html.H5([
            html.B("A"), "gentic ", html.B("R"), "isk ", html.B("I"), "ntelligence for ", html.B("AML"),
        ], className="mb-1"),
        html.P("Powered by Aria via Ollama", className="text-muted small mb-3"),

        html.Hr(className="my-2"),

        # Data summary badges
        html.Div([
            html.Span("Dataset", className="fw-semibold d-block mb-1 small"),
            dbc.Badge(f"{_total:,} accounts",    color="primary",  className="me-1 mb-1"),
            dbc.Badge(f"{_biz_count:,} business",  color="secondary",className="me-1 mb-1"),
            dbc.Badge(f"{_ind_count:,} individual", color="secondary",className="me-1 mb-1"),
            dbc.Badge(f"{_alert_count:,} alerts",   color="warning",  className="me-1 mb-1"),
            dbc.Badge(f"{_fp_count:,} FP",          color="danger",   className="me-1 mb-1"),
        ], className="mb-3"),

        html.Hr(className="my-2"),

        html.Span("Demo Prompts", className="fw-semibold d-block mb-1 small"),
        dbc.Select(
            id="demo-prompt-dropdown",
            options=[
                {"label": "Select a prompt...", "value": ""},
                {"label": "── AML Policy & Concepts ──", "value": "_h1", "disabled": True},
                *[{"label": p, "value": p} for p in DEMO_PROMPTS],
                {"label": "── Threshold Tuning ──", "value": "_h2", "disabled": True},
                *[{"label": p, "value": p} for p in SUGGESTED_PROMPTS[:3]],
                {"label": "── Dynamic Segmentation ──", "value": "_h3", "disabled": True},
                *[{"label": p, "value": p} for p in SUGGESTED_PROMPTS[3:6]],
                {"label": "── Rule-Level Sweep ──", "value": "_h4", "disabled": True},
                *[{"label": p, "value": p} for p in SUGGESTED_PROMPTS[6:]],
            ],
            value=DEMO_PROMPTS[0],
            className="mb-3",
            style={"fontSize": "0.8rem"},
        ),

        html.Hr(className="my-2"),

        # Model info
        html.Div([
            html.Span("Model", className="fw-semibold d-block mb-1 small"),
            dbc.Badge(OLLAMA_MODEL, color="info", className="me-1"),
            dbc.Badge("Ollama", color="dark"),
        ]),

        html.Hr(className="my-2"),

        # Drill-down button — enabled after a 2D sweep
        dbc.Button(
            "Drill-down Last 2D Sweep",
            id="drilldown-btn",
            color="outline-secondary",
            size="sm",
            disabled=True,
            className="w-100 mb-2",
        ),

        # Segment drilldown — enabled after clustering
        dbc.Button(
            "Segment Customer Drilldown",
            id="treemap-drilldown-btn",
            color="outline-success",
            size="sm",
            disabled=False,
            className="w-100 mb-2",
        ),

        # Cluster scatter — enabled after clustering, hidden by default for AML analysts
        dbc.Button(
            "Cluster Scatter Plot",
            id="scatter-btn",
            color="outline-secondary",
            size="sm",
            disabled=True,
            className="w-100 mb-2",
        ),

        # SAR priority worklist
        dbc.Button(
            "SAR Priority Worklist",
            id="sar-worklist-btn",
            color="danger",
            size="sm",
            disabled=_sar_scores_df is None,
            className="w-100",
        ),
    ])
], className="h-100 overflow-auto", style={"fontSize": "0.85rem"})

_chat_panel = html.Div([
    # Toolbar
    html.Div([
        dbc.ButtonGroup([
            dbc.Button("New Chat",  id="new-chat-btn",  color="secondary", outline=True, size="sm"),
            dbc.Button("Save Chat", id="save-chat-btn", color="secondary", outline=True, size="sm"),
        ]),
    ], className="px-3 py-1 d-flex justify-content-end",
       style={"borderBottom": "1px solid #dee2e6", "background": "#f8f9fa"}),
    html.Div([
        ChatComponent(
            id="chat-component",
            messages=[],
            class_name="AML AI",
            assistant_bubble_style={"maxWidth": "100%", "width": "100%"},
            supported_input_file_types=[".pdf", ".docx", ".doc", ".txt"],
        ),
        dbc.Button(
            "Stop",
            id="stop-btn",
            color="danger",
            size="sm",
            outline=True,
            style={"position": "absolute", "bottom": "70px", "right": "24px", "zIndex": 999, "opacity": 0.7},
        ),
        # Thinking indicator — absolute overlay, shown mid-callback via set_props
        html.Div(
            [dbc.Spinner(size="sm", color="secondary", type="border"),
             html.Span("Thinking…", className="ms-2 small text-muted")],
            id="chat-thinking-indicator",
            style={
                "display": "none",
                "position": "absolute",
                "bottom": "110px",
                "left": "50%",
                "transform": "translateX(-50%)",
                "background": "rgba(255,255,255,0.92)",
                "padding": "5px 16px",
                "borderRadius": "20px",
                "border": "1px solid #dee2e6",
                "zIndex": 1001,
                "alignItems": "center",
                "gap": "6px",
                "whiteSpace": "nowrap",
            },
        ),
    ], id="chat-scroll-container", style={
        "flex": "1",
        "minHeight": "0",
        "overflow": "auto",
        "overscrollBehaviorY": "contain",
        "position": "relative",
    }),
    # Charts rendered here — kept out of the messages prop entirely so
    # React never reconciles large Plotly JSON on every keystroke.
    html.Div(id="charts-panel", className="px-3 pb-2", style={
        "overflow": "auto",
        "maxHeight": "35vh",
        "borderTop": "1px solid #dee2e6",
    }),
], style={
    "display": "flex",
    "flexDirection": "column",
    "height": "calc(100vh - 80px)",
})

_about_panel = dbc.Collapse(
    dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([
                html.H6("What this demo does", className="fw-bold mb-2"),
                html.Ul([
                    html.Li("Threshold tuning — sweep FP/FN trade-offs as alert thresholds change"),
                    html.Li("SAR backtest — test how many true SARs a threshold configuration catches"),
                    html.Li("2D rule sweep — optimize two parameters with an interactive heatmap"),
                    html.Li("Dynamic Segmentation — cluster customers into behavioral groups using K-Means"),
                    html.Li("AML policy Q&A — compliance questions answered from FFIEC, Wolfsberg, FinCEN docs"),
                ], className="mb-0", style={"fontSize": "0.85rem"}),
            ], width=6),
            dbc.Col([
                html.H6("How to use", className="fw-bold mb-2"),
                html.P("Click a suggested prompt on the left or type your own question. "
                       "The AI routes your request to the correct analytics tool, runs the computation, "
                       "and returns results with charts.", style={"fontSize": "0.85rem"}),
                html.H6("Tech stack", className="fw-bold mb-2"),
                html.P("Fine-tuned Qwen 2.5 7B · Ollama · Plotly Dash · ChromaDB · scikit-learn",
                       style={"fontSize": "0.85rem", "color": "#aaa"}),
            ], width=6),
        ]),
    ]), className="mb-2 border-info"),
    id="about-collapse",
    is_open=False,
)

app.layout = dbc.Container([
    # Header
    dbc.Row([
        dbc.Col(
            html.H4("ARIA — Threshold Tuning & Dynamic Segmentation",
                    className="text-center my-3 fw-bold"),
        ),
        dbc.Col(
            dbc.Button("About this demo", id="about-toggle", color="info",
                       outline=True, size="sm", className="my-3 float-end"),
            width="auto",
        ),
    ], align="center"),
    _about_panel,
    # Body
    dbc.Row([
        dbc.Col(_sidebar,    width=3, className="pe-2",
                style={"height": "calc(100vh - 80px)", "overflowY": "auto", "position": "sticky", "top": "0"}),
        dbc.Col(_chat_panel, width=9, className="ps-2"),
    ], className="g-0"),

    # Download anchor for chat export
    dcc.Download(id="chat-download"),
    # Store for pending prompt from sidebar buttons
    dcc.Store(id="pending-prompt", data=None),
    dcc.Store(id="scroll-dummy"),
    dcc.Store(id="last-2d-sweep-store", data=None),
    dcc.Store(id="drilldown-heatmap-click-store", data=None),
    dcc.Store(id="treemap-store", data=None),
    dcc.Store(id="scatter-store", data=None),
    dcc.Store(id="charts-panel-store", data=None),
    dcc.Store(id="cluster-result-store", data=""),
    dcc.Store(id="scatter-resize-dummy"),
    # One-shot interval to inject welcome message after page load
    dcc.Interval(id="welcome-interval", interval=300, max_intervals=1),

    # ── Drill-down offcanvas ──────────────────────────────────────────────────
    dbc.Offcanvas(
        id="drilldown-offcanvas",
        title="2D Sweep Drill-down — Click a cell to see customers",
        placement="end",
        is_open=False,
        style={"width": "55vw"},
        children=[
            html.Div(id="drilldown-heatmap", style={"overflowX": "auto"}),
            html.Hr(),
            html.Div(
                "Run a 2D sweep, then click a cell above to see the customer breakdown.",
                id="drilldown-table-container",
                style={"fontSize": "0.85rem", "overflowY": "auto", "maxHeight": "40vh"},
            ),
        ],
    ),
    # ── SAR priority worklist offcanvas ──────────────────────────────────────
    dbc.Offcanvas(
        id="sar-worklist-offcanvas",
        title="SAR Priority Worklist",
        placement="end",
        is_open=False,
        style={"width": "70vw"},
        children=[
            # Model metrics
            html.Div(id="sar-worklist-metrics", className="mb-3"),
            # Threshold slider
            html.Div([
                html.Label("Minimum SAR Score threshold:", className="fw-semibold small mb-1"),
                dcc.Slider(
                    id="sar-threshold-slider",
                    min=50, max=99, step=1, value=80,
                    marks={50: "50%", 60: "60%", 70: "70%", 80: "80%", 90: "90%", 99: "99%"},
                    tooltip={"placement": "bottom", "always_visible": True},
                ),
            ], className="mb-3"),
            html.Hr(),
            html.Div(id="sar-worklist-table"),
        ],
    ),

    # ── Segment drilldown offcanvas ───────────────────────────────────────────
    dbc.Offcanvas(
        id="treemap-offcanvas",
        title="Segment Customer Drilldown — Click any tile to see customers",
        placement="end",
        is_open=False,
        style={"width": "65vw"},
        children=[
            dcc.Graph(
                id="treemap-drilldown-graph",
                config={"responsive": True},
                style={"height": "400px"},
            ),
            html.Hr(),
            html.Div(
                "Click any tile above to see customers in that segment.",
                id="treemap-customer-table",
                style={"fontSize": "0.85rem", "overflowY": "auto", "maxHeight": "45vh"},
            ),
        ],
    ),
    # ── Cluster scatter offcanvas ─────────────────────────────────────────────
    dbc.Offcanvas(
        id="scatter-offcanvas",
        title="Cluster Scatter Plot — for explainability",
        placement="end",
        is_open=False,
        style={"width": "65vw"},
        children=[
            dcc.Graph(
                id="scatter-offcanvas-graph",
                config={"responsive": True},
                useResizeHandler=True,
                style={"height": "480px", "width": "100%"},
            ),
        ],
    ),
    # ── Network graph offcanvas ───────────────────────────────────────────────
    dbc.Offcanvas(
        id="network-graph-offcanvas",
        title="Transaction Network Graph",
        placement="end",
        is_open=False,
        style={"width": "72vw"},
        children=[
            html.Div(id="network-graph-customer-badge", className="mb-2"),
            dcc.Graph(
                id="network-graph-figure",
                config={"responsive": True},
                style={"height": "380px"},
            ),
            html.Hr(),
            dcc.Graph(
                id="network-feature-bar",
                config={"responsive": True},
                style={"height": "260px"},
            ),
        ],
    ),
], fluid=True, style={"minHeight": "100vh"})


# ── Auto-scroll chat to bottom whenever messages update ───────────────────────
app.clientside_callback(
    """
    function(messages) {
        setTimeout(function() {
            var el = document.getElementById('chat-scroll-container');
            if (el) { el.scrollTop = el.scrollHeight; }
        }, 150);
        return null;
    }
    """,
    Output("scroll-dummy", "data"),
    Input("chat-component", "messages"),
    prevent_initial_call=True,
)

# ── Callbacks ─────────────────────────────────────────────────────────────────

@callback(
    Output("about-collapse", "is_open"),
    Input("about-toggle", "n_clicks"),
    State("about-collapse", "is_open"),
    prevent_initial_call=True,
)
def toggle_about(n_clicks, is_open):
    return not is_open

@callback(
    Output("chat-component", "messages", allow_duplicate=True),
    Input("welcome-interval", "n_intervals"),
    prevent_initial_call=True,
)
def show_welcome(n):
    return INITIAL_MESSAGES



@callback(
    Output("pending-prompt", "data", allow_duplicate=True),
    Input("demo-prompt-dropdown", "value"),
    prevent_initial_call=True,
)
def queue_demo_prompt(value):
    if not value or value.startswith("_h"):
        return no_update
    return {"query": value, "ts": time.time()}


@server.route("/stop", methods=["POST"])
def _stop_route():
    """Direct Flask route — bypasses Dash callback queue so it fires during inference."""
    _agent_stop_event.set()
    return flask.jsonify({"ok": True})


app.clientside_callback(
    """
    function(n_clicks) {
        if (n_clicks && n_clicks > 0) {
            fetch('/stop', {method: 'POST'});
        }
        return 0;
    }
    """,
    Output("stop-btn", "n_clicks"),
    Input("stop-btn", "n_clicks"),
    prevent_initial_call=True,
)


# Clientside callback — fires in the browser instantly when a demo prompt is selected.
# Appends the user bubble and shows the thinking indicator without a server round-trip,
# so the user sees feedback before handle_chat even starts.
app.clientside_callback(
    """
    function(pending_prompt, messages) {
        if (!pending_prompt || !pending_prompt.query) {
            return [window.dash_clientside.no_update, {"display": "none"}];
        }
        var userMsg = {role: "user", content: pending_prompt.query};
        var MAX = 30;
        var newMsgs = (messages || []).concat([userMsg]).slice(-MAX);
        var indicatorStyle = {
            display: "flex", position: "absolute", bottom: "110px",
            left: "50%", transform: "translateX(-50%)",
            background: "rgba(255,255,255,0.92)", padding: "5px 16px",
            borderRadius: "20px", border: "1px solid #dee2e6",
            zIndex: 1001, alignItems: "center", gap: "6px", whiteSpace: "nowrap"
        };
        return [newMsgs, indicatorStyle];
    }
    """,
    Output("chat-component", "messages", allow_duplicate=True),
    Output("chat-thinking-indicator", "style", allow_duplicate=True),
    Input("pending-prompt", "data"),
    State("chat-component", "messages"),
    prevent_initial_call=True,
)


@callback(
    Output("chat-component", "messages"),
    Output("last-2d-sweep-store", "data"),
    Output("treemap-store", "data"),
    Output("scatter-store", "data"),
    Output("charts-panel-store", "data"),
    Output("chat-thinking-indicator", "style", allow_duplicate=True),
    Input("chat-component", "new_message"),
    Input("pending-prompt", "data"),
    State("chat-component", "messages"),
    prevent_initial_call=True,
)
def handle_chat(new_message, pending_prompt, messages):
    import threading as _thr
    from dash import set_props as _set_props
    from agents.base_agent import stop_event
    stop_event.clear()             # clear any lingering stop signal from New Chat
    ctx     = callback_context
    trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
    print(f"[handle_chat] enter thread={_thr.current_thread().ident} trigger={trigger} msgs={len(messages or [])}", flush=True)

    # Determine query source
    if "pending-prompt" in trigger and pending_prompt:
        query    = pending_prompt["query"]
        user_msg = {"role": "user", "content": query}
        # User bubble + indicator are shown by the clientside callback immediately.
        # handle_chat builds the response and hides the indicator via its Output.
    elif new_message and new_message.get("role") == "user":
        content = new_message.get("content", "")
        if isinstance(content, list):
            # Attachment upload — extract text part; list is processed fully below
            query = next((b.get("text", "") for b in content
                          if isinstance(b, dict) and b.get("type") == "text"), "")
        else:
            query = content
        user_msg   = new_message
    else:
        return no_update, no_update, no_update, no_update, no_update, no_update

    if not query.strip() and not (isinstance(new_message, dict) and isinstance(new_message.get("content"), list)):
        return no_update, no_update, no_update, no_update, no_update, no_update

    # Lock the dropdown while the agent runs (best-effort race guard).
    _set_props("demo-prompt-dropdown", {"disabled": True})

    updated = messages + [user_msg]

    # ── Attachment handling ───────────────────────────────────────────────────
    # new_message.content can be a list when a file is attached:
    # [{type:"text", text:"..."}, {type:"attachment", file:"data:...", fileName:"...", fileType:"..."}]
    content_list = new_message.get("content") if isinstance(new_message, dict) else None
    attachment = None
    text_from_content = None
    if isinstance(content_list, list):
        for block in content_list:
            if isinstance(block, dict):
                if block.get("type") == "attachment":
                    attachment = block
                elif block.get("type") == "text" and block.get("text", "").strip():
                    text_from_content = block["text"].strip()

    if attachment:
        file_name = attachment.get("fileName", "uploaded_file")
        file_type = attachment.get("fileType", "")
        file_data = attachment.get("file", "")
        try:
            n_chunks = upload_kb.ingest_upload(file_data, file_name, file_type)
            ingest_msg = (
                f"**'{file_name}'** has been indexed — {n_chunks} chunks stored. "
                f"You can now ask questions about it."
            )
        except Exception as e:
            ingest_msg = f"Could not index '{file_name}': {e}"
            # Upload failed — always return immediately, never fall through to agent
            bot_response = {"role": "assistant", "content": ingest_msg}
            _set_props("demo-prompt-dropdown", {"disabled": False})
            return (updated + [bot_response])[-_MAX_CHAT_MESSAGES:], no_update, no_update, no_update, no_update, {"display": "none"}

        # If the user also typed a question alongside the file, continue to answer it.
        # Otherwise just confirm and return.
        if text_from_content:
            query = text_from_content
            # Fall through to normal agent routing below
            ingest_confirmation = ingest_msg
        else:
            bot_response = {"role": "assistant", "content": ingest_msg}
            _set_props("demo-prompt-dropdown", {"disabled": False})
            return (updated + [bot_response])[-_MAX_CHAT_MESSAGES:], no_update, no_update, no_update, no_update, {"display": "none"}
    else:
        ingest_confirmation = None

    # ── Help intent interception (before hitting the model) ──────────────────
    _help_pattern = re.compile(
        r'\b(what can you (do|help|assist)|what (features?|capabilities|tools?|functions?) (do you|does this|are)|'
        r'help me|show me what|what is this tool|what does this (app|tool|system)|'
        r'^help$)\b',
        re.IGNORECASE
    )
    if _help_pattern.search(query):
        help_text = (
            "Here's what I can help you with:\n\n"
            "1. **Threshold Tuning** — FP/FN trade-off analysis across alert thresholds by segment and transaction feature\n"
            "2. **SAR Backtest** — see how many SARs a specific rule catches at different thresholds\n"
            "3. **2D Sweep** — optimize two rule parameters simultaneously\n"
            "4. **Customer Segmentation** — K-Means clustering by transaction behavior (Business or Individual)\n"
            "5. **SAR Priority Worklist** — ranked list of SAR candidates by propensity score (red button in the sidebar)\n"
            "6. **Transaction Network Graph** — click any worklist row to see a customer's money flow network\n"
            "7. **AML Policy Q&A** — compliance questions answered from the regulatory knowledge base\n\n"
            "Try a prompt like: *'Show FP/FN trade-off for Business customers'* or *'Run SAR backtest for Activity Deviation ACH rule'*"
        )
        bot_response = {"role": "assistant", "content": help_text}
        _set_props("demo-prompt-dropdown", {"disabled": False})
        return (updated + [bot_response])[-_MAX_CHAT_MESSAGES:], no_update, no_update, no_update, no_update, {"display": "none"}

    # Clear any stale stop signal from a previous button click before starting a new query
    _agent_stop_event.clear()

    try:
        global _last_cluster_result, _last_cluster_raw_stats
        _thread_local.current_query = query
        _thread_local.last_2d_state = {}
        def _msg_text(m):
            c = m.get("content", "")
            if isinstance(c, str):
                return c
            if isinstance(c, list):
                for b in c:
                    if isinstance(b, dict) and b.get("type") == "text":
                        return b.get("text") or b.get("content") or ""
            return ""
        last_assistant = next(
            (_msg_text(m) for m in reversed(messages)
             if m.get("role") == "assistant" and _msg_text(m)),
            ""
        )
        # Build conversation history: last 2 user+assistant pairs before current query.
        # Truncate long messages so prior treemap/chart text doesn't flood the context.
        _MAX_HIST_CHARS = 120
        _history_raw = [
            {"role": m["role"], "content": _msg_text(m)}
            for m in (messages[:-1] if messages and messages[-1].get("role") == "user" else messages)
            if m.get("role") in ("user", "assistant") and _msg_text(m)
        ]
        history = [
            {
                "role": m["role"],
                "content": m["content"][:_MAX_HIST_CHARS] + ("…" if len(m["content"]) > _MAX_HIST_CHARS else ""),
            }
            for m in _history_raw[-4:]  # last 4 = 2 user+assistant pairs
        ]
        # If the user uploaded a document alongside a question, force policy agent
        # so the question is answered from the KB regardless of query keywords.
        if ingest_confirmation:
            agent_text, chart_results = orchestrator.policy_agent.run(query, tool_executor)
        else:
            # Prefer raw pre-computed stats over model prose for multi-turn cluster follow-ups
            _cluster_ctx = _last_cluster_raw_stats or _last_cluster_result
            agent_text, chart_results = orchestrator.run(query, tool_executor, last_assistant, history, _cluster_ctx)
        # Persist full clustering result for multi-turn cluster follow-ups.
        _seg_cluster_tools = {"ds_cluster_analysis", "cluster_analysis"}
        if any(r[0] in _seg_cluster_tools for r in chart_results):
            _last_cluster_result = agent_text or ""
        print(f"[app] raw agent_text ({len(agent_text or '')} chars): {repr((agent_text or '')[:300])}", flush=True)
    except Exception as e:
        import traceback
        traceback.print_exc()
        err = str(e)
        if "connection" in err.lower() or "connect" in err.lower() or "refused" in err.lower():
            msg_text = f"Cannot reach the model at {OLLAMA_BASE_URL}. Make sure Ollama is running and try again."
        elif "404" in err or "not found" in err.lower():
            msg_text = f"Model '{OLLAMA_MODEL}' was not found in Ollama. Run `ollama list` to see what is loaded."
        elif "429" in err or "rate limit" in err.lower() or "overload" in err.lower():
            msg_text = "The model is currently overloaded. Please wait a moment and try again."
        elif "timeout" in err.lower() or "timed out" in err.lower():
            msg_text = "The model took too long to respond. Try a simpler query or check if Ollama is still running."
        else:
            msg_text = "Something went wrong — please try again. If the problem persists, restart Ollama."
        bot_response = {"role": "assistant", "content": msg_text}
        _set_props("demo-prompt-dropdown", {"disabled": False})
        return (updated + [bot_response])[-_MAX_CHAT_MESSAGES:], no_update, no_update, no_update, no_update, {"display": "none"}

    # Parse DISPLAY_CLUSTERS directive and filter charts if present
    # Only honour the directive if the user actually asked to filter clusters
    _filter_keywords = re.search(
        r'\b(show only|only cluster|highest risk|lowest|top \d|filter)\b',
        query, re.IGNORECASE
    )
    _dc_match = re.search(r'DISPLAY_CLUSTERS:\s*([\d,\s]+)', agent_text or "")
    if _dc_match and chart_results and _cluster_cache and _filter_keywords:
        filter_nums = [int(x.strip()) for x in _dc_match.group(1).split(',') if x.strip().isdigit()]
        if filter_nums:
            import plotly.graph_objects as _go  # noqa: already imported at top via px
            df_clustered  = _cluster_cache["df_clustered"]
            scatter_fig   = _cluster_cache["scatter_fig"]
            customer_type = _cluster_cache.get("customer_type", "All")
            ss_dims = _DS_DIMS
            # Filter scatter traces
            keep_names = {f"Cluster {c}" for c in filter_nums}
            filtered_scatter = _go.Figure(
                data=[t for t in scatter_fig.data if t.name in keep_names],
                layout=scatter_fig.layout,
            )
            # Filter treemap df to requested clusters (0-based internally)
            cluster_vals = sorted(df_clustered["cluster"].unique())
            keep_0based  = {cluster_vals[c - 1] for c in filter_nums if 1 <= c <= len(cluster_vals)}
            df_filtered  = df_clustered[df_clustered["cluster"].isin(keep_0based)].copy()
            treemap_fig  = lambda_ds_performance.smartseg_tree_dynamic(
                df_filtered, f"{customer_type} — clusters {filter_nums}", dims=_DS_DIMS, df_rule_sweep=DF_RULE_SWEEP
            )
            chart_results = [("ds_cluster_analysis",
                              {"customer_type": customer_type, "filter_clusters": filter_nums},
                              (filtered_scatter, treemap_fig))]
    # Strip DISPLAY_CLUSTERS line and PRE-COMPUTED ANALYSIS markers from displayed text
    agent_text = re.sub(r'<eos>', '', agent_text or "").strip()               # Gemma 4 leaks <eos> tokens (strip all)
    agent_text = re.sub(r'\bEnd chunk\b\s*', '', agent_text).strip()          # leaked training artifact
    agent_text = re.sub(r'^Tool result for [^:\n]+:\n?', '', agent_text, flags=re.MULTILINE).strip()  # model copies tool-msg prefix
    agent_text = re.sub(r'^The PRE-COMPUTED[^\n]*\n?', '', agent_text).strip()  # leaked instruction header
    agent_text = re.sub(r'\s*DISPLAY_CLUSTERS:[\d,\s]*', '', agent_text).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED ANALYSIS.*?===.*?===\s*END PRE-COMPUTED ANALYSIS\s*===\n?(?:\([^\n]*\)\n?)?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'PRE-COMPUTED ANALYSIS[:\s]*\n?', '', agent_text).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED SEGMENT STATS.*?===.*?===\s*END PRE-COMPUTED SEGMENT STATS\s*===\n?(?:\([^\n]*\)\n?)?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED CLUSTER STATS.*?===.*?===\s*END PRE-COMPUTED CLUSTER STATS\s*===\n?(?:\([^\n]*\)\n?)?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED CLUSTER RULE SUMMARY.*?===.*?===\s*END CLUSTER RULE SUMMARY\s*===\n?(?:\([^\n]*\)\n?)?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED CLUSTER THRESHOLD ANALYSIS[^\n]*===\n?', '', agent_text).strip()
    agent_text = re.sub(r'\n?===\s*END CLUSTER THRESHOLD ANALYSIS\s*===\n?', '\n\n', agent_text).strip()
    agent_text = re.sub(r'\bPRE-COMPUTED CLUSTER THRESHOLD ANALYSIS\b\s*', '', agent_text).strip()
    # Catch malformed END markers where leading === is missing (model drops it occasionally)
    agent_text = re.sub(r'^END PRE-COMPUTED[^\n]*===\s*\n?', '', agent_text, flags=re.MULTILINE).strip()
    agent_text = re.sub(r'^END CLUSTER[^\n]*===\s*\n?', '', agent_text, flags=re.MULTILINE).strip()
    agent_text = re.sub(r'^#+\s*AML Domain Insight\s*\n?', '', agent_text, flags=re.MULTILINE | re.IGNORECASE).strip()
    agent_text = re.sub(r'\*\*AML Domain Insight\*\*\s*\n?', '', agent_text, flags=re.IGNORECASE).strip()
    # Ensure insight sentence after ADAPTIVE SUMMARY gets a paragraph break
    agent_text = re.sub(r'(- Net change:[^\n]+)\n([A-Z])', r'\1\n\n\2', agent_text)
    # Replace LaTeX math operators — model occasionally emits $\ge$/$\le$ instead of >= / <=
    agent_text = re.sub(r'\$\\geq\$', '≥', agent_text)
    agent_text = re.sub(r'\$\\leq\$', '≤', agent_text)
    agent_text = re.sub(r'\$\\ge\$',  '≥', agent_text)
    agent_text = re.sub(r'\$\\le\$',  '≤', agent_text)
    agent_text = re.sub(r'\$\\gt\$',  '>',  agent_text)
    agent_text = re.sub(r'\$\\lt\$',  '<',  agent_text)
    # Strip PRE-COMPUTED blocks — with === markers (primary) then without (fallback)
    agent_text = re.sub(r'===.*?PRE-COMPUTED SAR BACKTEST.*?===.*?===\s*END PRE-COMPUTED SAR BACKTEST\s*===\n?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'PRE-COMPUTED SAR BACKTEST.*?(?:===\s*END PRE-COMPUTED SAR BACKTEST\s*===\s*|END PRE-COMPUTED SAR BACKTEST\s*===\s*)', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED RULE SWEEP.*?===.*?===\s*END RULE SWEEP\s*===\n?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'PRE-COMPUTED RULE SWEEP.*?(?:===\s*END RULE SWEEP\s*===\s*|END RULE SWEEP\s*===\s*)', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'(?:===\s*)?END RULE SWEEP\s*===\s*', '', agent_text).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED RULE LIST.*?===.*?===\s*END RULE LIST\s*===\n?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'Available AML rules with SAR/FP performance.*?END RULE LIST\s*===\s*', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'(?:===\s*)?END RULE LIST\s*===\s*', '', agent_text).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED SEGMENT STATS.*?===.*?===\s*END PRE-COMPUTED SEGMENT STATS\s*===\n?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'PRE-COMPUTED SEGMENT STATS.*?(?:===\s*END PRE-COMPUTED SEGMENT STATS\s*===\s*|END PRE-COMPUTED SEGMENT STATS\s*===\s*)', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'===.*?PRE-COMPUTED 2D SWEEP.*?===.*?===\s*END 2D SWEEP\s*===\n?', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'PRE-COMPUTED 2D SWEEP.*?(?:===\s*END 2D SWEEP\s*===\s*|END 2D SWEEP\s*===\s*)', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'2D SWEEP[^\n]*Copy this verbatim[^\n]*\n?', '', agent_text, flags=re.IGNORECASE).strip()
    # Catch any remaining PRE-COMPUTED header lines (single-line remnants)
    agent_text = re.sub(r'^PRE-COMPUTED [^\n]+\n?', '', agent_text, flags=re.MULTILINE).strip()
    # Strip leaked instruction phrases and model meta-complaints
    agent_text = re.sub(r'^(Displaying|Tool call successful)[^\n]*\n?', '', agent_text, flags=re.MULTILINE | re.IGNORECASE).strip()
    # Strip stale "See chart/table below." that old training examples taught the model to emit
    agent_text = re.sub(r'\bSee chart/table below\.\s*', '', agent_text, flags=re.IGNORECASE).strip()
    agent_text = re.sub(r'\*?\(Detailed sweep (?:chart|table)[^)]*\)\*?\s*', '', agent_text, flags=re.IGNORECASE).strip()
    agent_text = re.sub(r'\[/PRE-COMPUTED RESULTS\]\s*', '', agent_text).strip()
    agent_text = re.sub(r'\[/Tool Output\]\s*', '', agent_text).strip()
    agent_text = re.sub(r'\[Tool Output\]\s*', '', agent_text).strip()
    agent_text = re.sub(r'\b(chart|heatmap|table|graph)\s+(shown|displayed)\s+above\b', r'\1 shown below', agent_text, flags=re.IGNORECASE)
    agent_text = re.sub(r'^NOTE: Data is simulated[^\n]*\n?', '', agent_text, flags=re.MULTILINE | re.IGNORECASE).strip()
    agent_text = re.sub(r'^The tool output is incomplete.*?verbatim\.\s*', '', agent_text, flags=re.DOTALL | re.IGNORECASE).strip()
    # Strip cluster report decorative lines and section markers
    agent_text = re.sub(r'^=+\s*(AML DYNAMIC SEGMENTATION REPORT|End Cluster Analysis)\s*\n?', '', agent_text, flags=re.MULTILINE).strip()
    agent_text = re.sub(r'^Cluster Summary\s*[—-]\s*\w+\s*\n?', '', agent_text, flags=re.MULTILINE).strip()
    agent_text = re.sub(r'^=+\s*\n?', '', agent_text, flags=re.MULTILINE).strip()
    # Strip Gemma 4 self-review chain-of-thought block
    agent_text = re.sub(r'\(Self-Correction.*?The response aligns with all rules\.\s*', '', agent_text, flags=re.DOTALL).strip()
    agent_text = re.sub(r'\(Self-Correction.*?\)\s*', '', agent_text, flags=re.DOTALL).strip()
    # Gemma 4 separator tokens: <channel|...> and <tool_call|> mark the boundary
    # between internal reasoning / tool-echo and the actual output.
    # Keep only what comes AFTER the last such token.
    _sep_match = list(re.finditer(r'<(?:channel|tool_call)\|[^>]*>', agent_text))
    if not _sep_match:
        # Also check bare <tool_call|> with no closing >
        _sep_match = list(re.finditer(r'<(?:channel|tool_call)\|>', agent_text))
    if _sep_match:
        agent_text = agent_text[_sep_match[-1].end():].strip()
    # Strip Gemma 4 string-delimiter tokens: <|"|>
    agent_text = re.sub(r'<\|"\|>', '', agent_text).strip()
    # Strip any remaining raw tool-call echo: tool_name{value:...} prefix
    agent_text = re.sub(r'^\w+\{value:.*?\}\s*', '', agent_text, flags=re.DOTALL).strip()
    # Strip self-generated JSON tool result blocks: {"risk_factor": ..., "pre_computed": ...}
    agent_text = re.sub(r'^\{[^{}]*"(?:risk_factor|pre_computed|sweep_param)[^{}]*\}', '', agent_text, flags=re.DOTALL).strip()
    # Strip leading punctuation artifacts (e.g. stray ] or ] \n left by token cleanup)
    agent_text = re.sub(r'^[\]\[)\s]+', '', agent_text).strip()

    # list_rules: strip the ### header (chart replaces it), prepend standard header,
    # and trust the model's insight sentence. Fall back server-side only if model returned empty.
    if any(name == "list_rules" for name, _, _ in chart_results):
        # Strip any ### header line the model emitted (we prepend our own)
        agent_text = re.sub(r'^###[^\n]*\n+', '', (agent_text or "")).strip()
        if not agent_text:
            _rule_num_match = re.search(r'\brule\s+(\d+)\b', (query or "").lower())
            if _rule_num_match and DF_RULE_SWEEP is not None:
                rule_num = _rule_num_match.group(1)
                rule_names = sorted(DF_RULE_SWEEP["risk_factor"].unique().tolist())
                names_str = ", ".join(rule_names)
                agent_text = (
                    f"There is no rule called 'rule {rule_num}' — rules are identified by name, not number. "
                    f"The {len(rule_names)} available rules are: {names_str}. "
                    f"Please specify the rule by name."
                )
            else:
                insight = "Rule performance summary — detailed table shown below."
                if DF_RULE_SWEEP is not None:
                    fp_by_rule = (
                        DF_RULE_SWEEP[DF_RULE_SWEEP["is_sar"] == 0]
                        .groupby("risk_factor").size().sort_values(ascending=False)
                    )
                    top3 = fp_by_rule.head(3)
                    if len(top3):
                        parts = [f"{rf} ({fp:,} FP)" for rf, fp in top3.items()]
                        insight += "\n\nRules with the most false positives: " + ", ".join(parts) + "."
                agent_text = insight
        else:
            agent_text = f"Rule performance summary — detailed table shown below.\n\n{agent_text}"

    # If model returned empty/whitespace after a tool call, generate a
    # data-driven insight server-side rather than a generic fallback.
    if chart_results and not (agent_text or "").strip():
        first_tool, first_input, _ = chart_results[0]

        if first_tool == "rule_sar_backtest" and DF_RULE_SWEEP is not None:
            rf      = first_input.get("risk_factor", first_input.get("rule_code", ""))
            sp      = first_input.get("sweep_param", None)
            cluster = first_input.get("cluster", None)
            df_rb   = _filter_by_cluster(DF_RULE_SWEEP, cluster)
            df_rule = df_rb[df_rb["risk_factor"].str.lower() == rf.lower()] if rf else df_rb
            if sp and sp in df_rule.columns and len(df_rule):
                total_sar  = int(df_rule["is_sar"].sum())
                thresholds = sorted(df_rule[sp].dropna().unique())
                rows = []
                for t in thresholds[:6]:
                    caught = int(df_rule[(df_rule[sp] <= t) & (df_rule["is_sar"] == 1)].shape[0])
                    catch_pct = round(100 * caught / total_sar, 1) if total_sar else 0
                    rows.append(f"  {sp}={t}: **{caught}** SARs caught ({catch_pct}%)")
                cluster_tag = f" — Cluster {cluster}" if cluster else ""
                agent_text = (
                    f"**{rf} / {sp}{cluster_tag}** — {total_sar} total SARs in scope.\n\n"
                    + "\n".join(rows)
                )
            else:
                agent_text = f"SAR backtest complete — see chart below for **{rf}**."

        elif first_tool == "threshold_tuning" and DF is not None:
            seg = first_input.get("segment", "")
            col = first_input.get("threshold_column", "")
            agent_text = f"Threshold sweep for **{seg}** customers by **{col}** — chart below."

        elif first_tool in ("sar_backtest",):
            seg = first_input.get("segment", "")
            col = first_input.get("threshold_column", "")
            agent_text = f"SAR backtest for **{seg}** / **{col}** — chart below."

        elif first_tool == "rule_2d_sweep":
            rf = first_input.get("risk_factor", "")
            agent_text = f"2D sweep for **{rf}** — use the **Drill-down** button in the sidebar for interactive exploration."

        elif first_tool in ("ds_cluster_analysis", "cluster_analysis"):
            ct = first_input.get("customer_type", "All")
            # Try to answer attribute questions from the stats block before falling back to generic text
            _attr_map = [
                (["age", "account age", "oldest", "youngest", "tenure"], "Account Age (years)"),
                (["income"],                                               "Income"),
                (["balance"],                                              "Current Balance"),
                (["transaction amount", "avg transaction amount"],         "Avg Transaction Amount"),
                (["transaction volume", "monthly volume", "monthly transaction"], "Monthly Transaction Volume"),
                (["weekly transaction", "avg weekly", "transactions per week"], "Avg Weekly Transactions"),
            ]
            q_l = (query or "").lower()
            target_attr = next((lbl for kws, lbl in _attr_map if any(kw in q_l for kw in kws)), None)
            stats_text  = _last_cluster_raw_stats or ""
            parsed = {}
            if target_attr and stats_text:
                cur = None
                for line in stats_text.split("\n"):
                    cm = re.search(r'\*\*Cluster (\d+)\*\*', line)
                    if cm:
                        cur = int(cm.group(1))
                    elif cur and target_attr in line:
                        vm = re.search(r'\*\*\$?([\d,]+\.?\d*)\*\*', line)
                        if vm:
                            parsed[cur] = float(vm.group(1).replace(",", ""))
            if parsed:
                _hi = any(w in q_l for w in ["highest", "most", "oldest", "largest", "max", "best", "top"])
                _lo = any(w in q_l for w in ["lowest", "least", "youngest", "smallest", "min", "worst"])
                if _hi:
                    c = max(parsed, key=parsed.get)
                    agent_text = f"Cluster {c} has the highest {target_attr}: **{parsed[c]:,.1f}**."
                elif _lo:
                    c = min(parsed, key=parsed.get)
                    agent_text = f"Cluster {c} has the lowest {target_attr}: **{parsed[c]:,.1f}**."
                else:
                    parts = [f"Cluster {c}: {v:,.1f}" for c, v in sorted(parsed.items())]
                    agent_text = f"{target_attr} by cluster: " + ", ".join(parts) + "."
            else:
                agent_text = (
                    f"Clustering complete for **{ct}** customers. "
                    "Use **Segment Customer Drilldown** and **Cluster Scatter Plot** in the left sidebar to explore."
                )

        else:
            agent_text = "Results shown below."

    prefix = (ingest_confirmation + "\n\n") if ingest_confirmation else ""
    all_chart_dicts = []
    if chart_results:
        content = [{"type": "text", "text": prefix + agent_text}] if (prefix + agent_text) else []
        for tool_name, tool_input, fig in chart_results:
            text_blocks, new_charts = _chart_content(tool_name, tool_input, fig)
            content.extend(text_blocks)
            all_chart_dicts.extend(new_charts)
        bot_response = {"role": "assistant", "content": content}
    else:
        bot_response = {"role": "assistant", "content": prefix + (agent_text or "(No response)")}

    # Signal drill-down store if a 2D sweep was just run; store full state for per-user drill-down
    sweep_store = no_update
    _local_2d = getattr(_thread_local, 'last_2d_state', {})
    if _local_2d and any(tn == "rule_2d_sweep" for tn, _, _ in (chart_results or [])):
        sweep_store = _json_safe(_local_2d)

    # Signal treemap store if a clustering tool just ran
    treemap_store = no_update
    if any(tn in ("cluster_analysis", "ds_cluster_analysis") for tn, _, _ in (chart_results or [])):
        treemap_store = {"ts": time.time()}

    # Extract scatter figure for the scatter offcanvas button
    scatter_store = no_update
    for tn, _, f in (chart_results or []):
        if tn in ("cluster_analysis", "ds_cluster_analysis") and isinstance(f, tuple):
            # scatter is always second-to-last in the tuple: (stats?, scatter, treemap)
            scatter_fig = f[-2]
            if scatter_fig is not None:
                # Explicit height required — offcanvas is display:none when store updates,
                # so responsive mode can't infer container size and renders blank without it.
                scatter_fig.update_layout(height=480)
                scatter_store = scatter_fig.to_dict()
            break

    _set_props("demo-prompt-dropdown", {"disabled": False})
    final_messages = (updated + [bot_response])[-_MAX_CHAT_MESSAGES:]
    charts_out = all_chart_dicts if all_chart_dicts else no_update
    return final_messages, sweep_store, treemap_store, scatter_store, charts_out, {"display": "none"}


@callback(
    Output("charts-panel", "children"),
    Input("charts-panel-store", "data"),
)
def render_charts_panel(charts):
    if not charts:
        return []
    children = [html.Hr(className="my-3")]
    for c in charts:
        h = c.get("height", 460)
        children.append(
            html.H6(
                c["title"],
                className="fw-semibold text-muted mt-3 mb-1",
                style={"fontSize": "0.85rem"},
            )
        )
        children.append(
            dcc.Graph(
                figure=c["figure"],
                config={"responsive": True},
                style={"height": f"{h}px"},
            )
        )
    return children


# ── Drill-down callbacks ──────────────────────────────────────────────────────

@callback(
    Output("drilldown-heatmap", "children"),
    Output("drilldown-btn", "disabled"),
    Input("last-2d-sweep-store", "data"),
)
def refresh_drilldown_heatmap(store_data):
    """Render CSS-based heatmap — browser-agnostic, no canvas/WebGL."""
    if not store_data or not store_data.get("grid"):
        return "", True
    html_table = make_figures.rule_2d_heatmap_html(store_data["grid"])
    return html_table, False


@callback(
    Output("drilldown-offcanvas", "is_open"),
    Input("drilldown-btn", "n_clicks"),
    State("drilldown-offcanvas", "is_open"),
    prevent_initial_call=True,
)
def toggle_drilldown(n, is_open):
    return not is_open


@callback(
    Output("drilldown-heatmap-click-store", "data"),
    Input({"type": "heatmap-cell", "p1": ALL, "p2": ALL}, "n_clicks"),
    prevent_initial_call=True,
)
def heatmap_cell_clicked(n_clicks_list):
    """Convert pattern-matched cell click into a (p1_idx, p2_idx) store entry."""
    ctx = callback_context
    if not ctx.triggered or not any(n_clicks_list):
        return no_update
    triggered_id = ctx.triggered[0]["prop_id"].split(".")[0]
    import json
    cell_id = json.loads(triggered_id)
    return {"p1_idx": cell_id["p1"], "p2_idx": cell_id["p2"]}


@callback(
    Output("drilldown-table-container", "children"),
    Input("drilldown-heatmap-click-store", "data"),
    State("last-2d-sweep-store", "data"),
    prevent_initial_call=True,
)
def drilldown_on_click(click_data, sweep_state):
    """Filter customers at the clicked (p1_idx, p2_idx) cell."""
    if not click_data or not sweep_state or not sweep_state.get("grid"):
        return "Click a cell in the heatmap to see customer breakdown."

    grid   = sweep_state["grid"]
    p2_val = grid["p2_vals"][int(click_data["p2_idx"])]
    p1_val = grid["p1_vals"][int(click_data["p1_idx"])]

    rf      = sweep_state["risk_factor"]
    param1  = sweep_state["param1"]
    param2  = sweep_state["param2"]
    cluster = sweep_state.get("cluster")

    df_sweep = _filter_by_cluster(DF_RULE_SWEEP, cluster)
    tp_df, fp_df, fn_df, tn_df, col1, col2 = lambda_rule_analysis.compute_2d_drilldown(
        df_sweep, rf, param1, p1_val, param2, p2_val
    )

    if tp_df is None:
        return "Could not compute drill-down — rule data unavailable."

    grid    = sweep_state["grid"]
    p1_lbl  = grid["p1_label"]
    p2_lbl  = grid["p2_label"]

    def _fmt(v):
        try:
            return f"{float(v):,.1f}"
        except Exception:
            return str(v)

    def _make_table(df, status_label, status_color):
        if df is None or len(df) == 0:
            return html.P(f"No {status_label} customers at this cell.", className="text-muted small mb-1")
        cols = ["customer_id", "customer_type"]
        if col1 in df.columns: cols.append(col1)
        if col2 in df.columns and col2 != col1: cols.append(col2)
        display_df = df[cols].copy().head(50)
        for c in [col1, col2]:
            if c in display_df.columns:
                display_df[c] = display_df[c].apply(_fmt)
        return html.Div([
            html.P(
                f"{status_label} — {len(df)} customer(s)"
                + (" (showing first 50)" if len(df) > 50 else ""),
                className=f"fw-semibold text-{status_color} mb-1 small",
            ),
            dash_table.DataTable(
                data=display_df.to_dict("records"),
                columns=[{"name": c, "id": c} for c in display_df.columns],
                style_table={"overflowX": "auto", "marginBottom": "12px"},
                style_cell={"fontSize": "0.78rem", "padding": "3px 6px", "textAlign": "left"},
                style_header={"fontWeight": "bold", "backgroundColor": "#f8f9fa"},
                page_size=20,
            ),
        ])

    header = html.Div([
        html.P(
            f"{rf} @ {param1}={_fmt(p1_val)} ({p1_lbl}), {param2}={_fmt(p2_val)} ({p2_lbl})",
            className="fw-bold mb-2",
        ),
        html.Div([
            dbc.Badge(f"TP={len(tp_df)}", color="success", className="me-1"),
            dbc.Badge(f"FP={len(fp_df)}", color="warning",  className="me-1"),
            dbc.Badge(f"FN={len(fn_df)}", color="danger",   className="me-1"),
            dbc.Badge(f"TN={len(tn_df)}", color="secondary"),
        ], className="mb-3"),
    ])

    return html.Div([
        header,
        _make_table(fn_df, "Missed SARs (FN)", "danger"),
        _make_table(fp_df, "False Positives (FP)", "warning"),
        _make_table(tp_df, "Caught SARs (TP)", "success"),
    ])


# ── Segment drilldown callbacks ───────────────────────────────────────────────

@callback(
    Output("treemap-drilldown-graph", "figure"),
    Input("treemap-store", "data"),
)
def refresh_treemap_offcanvas(store_data):
    """Re-render the treemap in the offcanvas whenever clustering runs."""
    df_clustered = _cluster_cache.get("df_clustered")
    if df_clustered is None:
        return {}
    ct = _cluster_cache.get("customer_type", "All")
    fig = lambda_ds_performance.smartseg_tree_dynamic(
        df_clustered, ct, dims=_DS_DIMS, df_rule_sweep=DF_RULE_SWEEP
    )
    fig.update_layout(height=380, margin=dict(t=30, b=10, l=10, r=10))
    return fig


@callback(
    Output("treemap-offcanvas", "is_open"),
    Input("treemap-drilldown-btn", "n_clicks"),
    State("treemap-offcanvas", "is_open"),
    prevent_initial_call=True,
)
def toggle_treemap_offcanvas(n, is_open):
    return not is_open


@callback(
    Output("scatter-offcanvas-graph", "figure"),
    Output("scatter-btn", "disabled"),
    Input("scatter-store", "data"),
)
def refresh_scatter_offcanvas(scatter_data):
    """Populate the scatter offcanvas and enable its button whenever clustering runs."""
    if not scatter_data:
        return {}, True
    return scatter_data, False


@callback(
    Output("scatter-offcanvas", "is_open"),
    Input("scatter-btn", "n_clicks"),
    State("scatter-offcanvas", "is_open"),
    prevent_initial_call=True,
)
def toggle_scatter_offcanvas(n, is_open):
    return not is_open


# Firefox fix: call Plotly.Plots.resize() directly on the graph element after
# the offcanvas animation completes.  window.dispatchEvent('resize') is ignored
# by Firefox when the element was not visible at layout time.
app.clientside_callback(
    """
    function(is_open) {
        if (is_open) {
            setTimeout(function() {
                var el = document.getElementById('scatter-offcanvas-graph');
                if (el && window.Plotly) {
                    Plotly.Plots.resize(el);
                } else {
                    window.dispatchEvent(new Event('resize'));
                }
            }, 400);
        }
        return window.dash_clientside.no_update;
    }
    """,
    Output("scatter-resize-dummy", "data"),
    Input("scatter-offcanvas", "is_open"),
)


@callback(
    Output("treemap-customer-table", "children"),
    Input("treemap-drilldown-graph", "clickData"),
    prevent_initial_call=True,
)
def treemap_customer_drilldown(click_data):
    """Show customer table for the clicked treemap tile."""
    if not click_data or not _cluster_cache.get("df_enriched") is not None:
        return "Click any tile above to see customers in that segment."

    node_id = click_data["points"][0].get("id", "")
    df = _filter_treemap_node(node_id, _cluster_cache["df_enriched"])

    if df is None or len(df) == 0:
        return html.P("No customers found for this segment.", className="text-muted small")

    # Join SAR scores if available
    if _sar_scores_df is not None and "customer_id" in df.columns:
        sar_map = _sar_scores_df.set_index("customer_id")["sar_prob"]
        df = df.copy()
        df["sar_prob"] = df["customer_id"].map(sar_map)

    # Build display columns
    id_col   = next((c for c in ("customer_id", "CUSTOMER_ID", "cust_id") if c in df.columns), None)
    show_cols = []
    if id_col:
        show_cols.append(id_col)
    for c in ("customer_type", "cluster_label"):
        if c in df.columns:
            show_cols.append(c)
    if "sar_prob" in df.columns:
        show_cols.append("sar_prob")
    for c in ("is_alerted", "is_fp", "is_fn", "is_sar"):
        if c in df.columns:
            show_cols.append(c)

    display_df = df[show_cols].copy().head(500)
    rename_map = {
        "is_alerted":  "Alerted",
        "is_fp":       "False Positive",
        "is_fn":       "False Negative",
        "is_sar":      "SAR",
        "sar_prob":    "SAR Score",
        "cluster_label": "Cluster",
        "customer_type": "Type",
    }
    display_df = display_df.rename(columns=rename_map)
    for col in ("Alerted", "False Positive", "False Negative", "SAR"):
        if col in display_df.columns:
            display_df[col] = display_df[col].map({1: "Yes", 0: "No"})
    if "SAR Score" in display_df.columns:
        display_df["SAR Score"] = (display_df["SAR Score"] * 100).round(1).astype(str) + "%"

    n_sar = int(df["is_sar"].sum())   if "is_sar"     in df.columns else 0
    n_fp  = int(df["is_fp"].sum())    if "is_fp"      in df.columns else 0
    n_fn  = int(df["is_fn"].sum())    if "is_fn"      in df.columns else 0
    n_alt = int(df["is_alerted"].sum()) if "is_alerted" in df.columns else 0

    header = html.Div([
        html.P(
            f"Segment: {node_id or 'All'} — {len(df):,} customers"
            + (f" (showing first 500)" if len(df) > 500 else ""),
            className="fw-bold mb-1 small",
        ),
        html.Div([
            dbc.Badge(f"Alerted: {n_alt}", color="warning",   className="me-1"),
            dbc.Badge(f"FP: {n_fp}",       color="danger",    className="me-1"),
            dbc.Badge(f"FN: {n_fn}",       color="secondary", className="me-1"),
            dbc.Badge(f"SAR: {n_sar}",     color="success",   className="me-1"),
        ], className="mb-2"),
    ])

    table = dash_table.DataTable(
        data=display_df.to_dict("records"),
        columns=[{"name": c, "id": c} for c in display_df.columns],
        style_table={"overflowX": "auto", "overflowY": "auto", "maxHeight": "38vh"},
        style_cell={"fontSize": "0.78rem", "padding": "3px 6px", "textAlign": "left"},
        style_header={"fontWeight": "bold", "backgroundColor": "#f8f9fa"},
        filter_action="native",
        sort_action="native",
        page_size=50,
    )

    return html.Div([header, table])


# ── SAR worklist callbacks ────────────────────────────────────────────────────

@callback(
    Output("sar-worklist-offcanvas", "is_open"),
    Input("sar-worklist-btn", "n_clicks"),
    State("sar-worklist-offcanvas", "is_open"),
    prevent_initial_call=True,
)
def toggle_sar_worklist(n, is_open):
    return not is_open


@callback(
    Output("sar-worklist-metrics", "children"),
    Output("sar-worklist-table", "children"),
    Input("sar-threshold-slider", "value"),
)
def update_sar_worklist(threshold_pct):
    if _sar_scores_df is None:
        return html.P("SAR scorer not available.", className="text-muted"), ""

    threshold = threshold_pct / 100.0
    df        = _sar_scores_df.copy()

    # Join cluster info if available
    if _cluster_cache.get("df_enriched") is not None and "customer_id" in df.columns:
        cl_df = _cluster_cache["df_enriched"][["customer_id", "cluster_label"]].drop_duplicates("customer_id")
        df    = df.merge(cl_df, on="customer_id", how="left")

    total_alerts  = len(df)
    above_thresh  = df[df["sar_prob"] >= threshold]
    n_above       = len(above_thresh)
    reduction_pct = round(100 * (1 - n_above / total_alerts), 1) if total_alerts > 0 else 0
    n_sar_caught  = int(above_thresh["is_sar"].sum()) if "is_sar" in above_thresh.columns else 0
    n_sar_total   = int(df["is_sar"].sum())           if "is_sar" in df.columns else 0
    sar_recall    = round(100 * n_sar_caught / n_sar_total, 1) if n_sar_total > 0 else 0

    metrics = html.Div([
        dbc.Row([
            dbc.Col(dbc.Card(dbc.CardBody([
                html.P("Model ROC-AUC", className="text-muted small mb-1"),
                html.H5(f"{_sar_roc_auc:.3f}", className="fw-bold text-info mb-0"),
            ]), className="text-center"), width=3),
            dbc.Col(dbc.Card(dbc.CardBody([
                html.P("Alerts to Investigate", className="text-muted small mb-1"),
                html.H5(f"{n_above:,} / {total_alerts:,}", className="fw-bold text-warning mb-0"),
            ]), className="text-center"), width=3),
            dbc.Col(dbc.Card(dbc.CardBody([
                html.P("Workload Reduction", className="text-muted small mb-1"),
                html.H5(f"{reduction_pct}%", className="fw-bold text-success mb-0"),
            ]), className="text-center"), width=3),
            dbc.Col(dbc.Card(dbc.CardBody([
                html.P("SAR Recall", className="text-muted small mb-1"),
                html.H5(f"{sar_recall}%", className="fw-bold text-danger mb-0"),
            ]), className="text-center"), width=3),
        ], className="g-2"),
    ])

    # Build display table
    show_cols = []
    for c in ("customer_id", "customer_type", "cluster_label"):
        if c in above_thresh.columns: show_cols.append(c)
    show_cols.append("sar_prob")
    for c in ("alert_count", "rule_count", "max_z_score", "avg_weekly_trxn_amt", "trxn_amt_monthly", "is_sar"):
        if c in above_thresh.columns: show_cols.append(c)

    display_df = above_thresh[show_cols].copy().head(1000)
    display_df["sar_prob"] = (display_df["sar_prob"] * 100).round(1).astype(str) + "%"
    rename = {
        "sar_prob":           "SAR Score",
        "customer_type":      "Type",
        "cluster_label":      "Cluster",
        "alert_count":        "Alerts",
        "rule_count":         "Rules Triggered",
        "max_z_score":        "Max Z-Score",
        "avg_weekly_trxn_amt":"Avg Weekly Amt",
        "trxn_amt_monthly":   "Monthly Amt",
        "is_sar":             "Known SAR",
        "customer_id":        "Customer ID",
    }
    display_df = display_df.rename(columns=rename)
    if "Known SAR" in display_df.columns:
        display_df["Known SAR"] = display_df["Known SAR"].map({1: "Yes", 0: "No"})
    for col in ("Avg Weekly Amt", "Monthly Amt"):
        if col in display_df.columns:
            display_df[col] = display_df[col].apply(lambda v: f"${v:,.0f}" if pd.notna(v) else "")

    COL_TOOLTIPS = {
        "Customer ID":    "Unique identifier for the customer account under review.",
        "Type":           "Customer segment: INDIVIDUAL or BUSINESS.",
        "Cluster":        "Behavioral cluster assigned by K-Means segmentation (C1, C2, …).",
        "SAR Score":      "SAR propensity score (0–100%) from the Random Forest model. "
                          "Higher = more likely to warrant a SAR filing. "
                          "Trained on alert, rule, transaction, and network graph features.",
        "Alerts":         "Total number of AML alerts generated for this customer across all rules.",
        "Rules Triggered":"Number of distinct monitoring rules that fired for this customer.",
        "Max Z-Score":    "Highest z-score across all alert rules that fired. "
                          "Measures how many standard deviations above the segment mean "
                          "the customer's behavior falls on the worst-scoring rule.",
        "Avg Weekly Amt": "Average per-transaction dollar amount (USD). "
                          "Reflects the typical size of a single transaction for this customer.",
        "Monthly Amt":    "Total transaction volume (USD) summed across all transactions in the monitored month. "
                          "Not an average — a high value means high overall activity.",
        "Known SAR":      "Ground-truth label: Yes = a SAR was actually filed for this customer "
                          "(used for model training and recall measurement).",
    }

    table = dash_table.DataTable(
        id="sar-worklist-datatable",
        data=display_df.to_dict("records"),
        columns=[{"name": c, "id": c} for c in display_df.columns],
        tooltip_header={c: COL_TOOLTIPS[c] for c in display_df.columns if c in COL_TOOLTIPS},
        tooltip_delay=0,
        tooltip_duration=None,
        style_table={"overflowX": "auto", "overflowY": "auto", "maxHeight": "55vh"},
        style_cell={"fontSize": "0.78rem", "padding": "3px 8px", "textAlign": "left"},
        style_header={"fontWeight": "bold", "backgroundColor": "#f8f9fa",
                      "textDecoration": "underline dotted", "cursor": "help"},
        style_data_conditional=[
            {"if": {"filter_query": '{Known SAR} = "Yes"'},
             "backgroundColor": "#fff3cd", "fontWeight": "bold"},
        ],
        filter_action="native",
        sort_action="native",
        sort_by=[{"column_id": "SAR Score", "direction": "desc"}],
        page_size=50,
    )

    return metrics, html.Div([
        html.P(
            f"Showing {min(n_above, 1000):,} of {n_above:,} alerts at or above {threshold_pct}% threshold"
            + (" (capped at 1,000 rows)" if n_above > 1000 else ""),
            className="text-muted small mb-2",
        ),
        table,
    ])


@callback(
    Output("network-graph-offcanvas",    "is_open"),
    Output("network-graph-figure",       "figure"),
    Output("network-feature-bar",        "figure"),
    Output("network-graph-customer-badge", "children"),
    Input("sar-worklist-datatable",      "active_cell"),
    State("sar-worklist-datatable",      "data"),
    State("network-graph-offcanvas",     "is_open"),
    prevent_initial_call=True,
)
def show_network_graph(active_cell, table_data, is_open):
    import plotly.graph_objects as go

    _empty = go.Figure(layout=go.Layout(
        paper_bgcolor="#1a1a1a", plot_bgcolor="#111",
        font=dict(color="#eee"),
        xaxis=dict(visible=False), yaxis=dict(visible=False),
        annotations=[dict(text="No data available", showarrow=False,
                          font=dict(size=14, color="#aaa"), xref="paper", yref="paper", x=0.5, y=0.5)],
    ))

    if not active_cell or not table_data:
        return no_update, no_update, no_update, no_update

    row      = table_data[active_cell["row"]]
    cid_key  = next((k for k in row if "customer" in k.lower() and "id" in k.lower()), None)
    if cid_key is None:
        return no_update, no_update, no_update, no_update

    customer_id = row[cid_key]
    sar_score   = row.get("SAR Score", "—")
    cust_type   = row.get("Type", "")
    cluster     = row.get("Cluster", "")
    known_sar   = row.get("Known SAR", "")

    badge = html.Div([
        dbc.Badge(f"Customer: {customer_id}", color="primary", className="me-2"),
        dbc.Badge(f"SAR Score: {sar_score}",  color="danger",  className="me-2"),
        dbc.Badge(cust_type,                  color="secondary", className="me-2") if cust_type else None,
        dbc.Badge(cluster,                    color="info",    className="me-2") if cluster else None,
        dbc.Badge("Known SAR" if known_sar == "Yes" else "Not filed",
                  color="warning" if known_sar == "Yes" else "success", className="me-2"),
        html.Small(" Click any row in the worklist to view its network graph.",
                   className="text-muted ms-2"),
    ], className="d-flex align-items-center flex-wrap")

    net_fig, feat_fig = build_network_graph(customer_id)

    if net_fig is None:
        net_fig = _empty
    if feat_fig is None:
        feat_fig = _empty

    return True, net_fig, feat_fig, badge


# ── New Chat ──────────────────────────────────────────────────────────────────
@callback(
    Output("chat-component", "messages", allow_duplicate=True),
    Output("cluster-result-store", "data", allow_duplicate=True),
    Output("pending-prompt", "data", allow_duplicate=True),
    Output("chat-thinking-indicator", "style", allow_duplicate=True),
    Input("new-chat-btn", "n_clicks"),
    prevent_initial_call=True,
)
def new_chat(_):
    from agents.base_agent import stop_event
    global _last_cluster_result, _last_cluster_raw_stats
    stop_event.set()               # signals in-flight generation to abort
    _last_cluster_result    = ""
    _last_cluster_raw_stats = ""
    return [], "", None, {"display": "none"}


# ── Save Chat ─────────────────────────────────────────────────────────────────
@callback(
    Output("chat-download", "data"),
    Input("save-chat-btn", "n_clicks"),
    State("chat-component", "messages"),
    prevent_initial_call=True,
)
def save_chat(_, messages):
    if not messages:
        return no_update
    docx_bytes = _messages_to_docx(messages)
    filename   = f"aria_session_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return dcc.send_bytes(docx_bytes, filename)


if __name__ == "__main__":
    app.run(debug=False, port=7860, host="0.0.0.0", use_reloader=False, threaded=True)
